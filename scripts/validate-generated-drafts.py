#!/usr/bin/env python3
"""Structural and distribution checks for generated PMD DOCX files."""

from __future__ import annotations

import argparse
import hashlib
import json
from pathlib import Path

from docx import Document


REQUIRED_HEADINGS = {
    "Contenido",
    "Síntesis ejecutiva",
    "1. Información general",
    "2. Diagnóstico municipal",
    "3. Fortalezas y debilidades basadas en el diagnóstico",
    "4. Visión Municipal",
    "5. Plan de acción a acordar por el CDM",
    "Fuentes y trazabilidad",
}

FORBIDDEN_TEXT = {
    "FODA",
    "Borrador técnico",
    "BORRADOR TÉCNICO",
    "3.2 Insumos técnicos para oportunidades y amenazas",
    "6. Seguimiento y evaluación",
}


def digest(path: Path) -> str:
    value = hashlib.sha256()
    with path.open("rb") as handle:
        for chunk in iter(lambda: handle.read(1024 * 1024), b""):
            value.update(chunk)
    return value.hexdigest()


def main() -> int:
    parser = argparse.ArgumentParser()
    parser.add_argument("--manifest", type=Path, required=True)
    parser.add_argument("--docx-dir", type=Path, required=True)
    parser.add_argument("--web-dir", type=Path, required=True)
    args = parser.parse_args()

    manifest = json.loads(args.manifest.read_text(encoding="utf-8"))
    rows = manifest["municipalities"]
    errors: list[str] = []
    stats = {
        "documents": len(rows),
        "min_paragraphs": None,
        "max_paragraphs": 0,
        "min_tables": None,
        "max_tables": 0,
        "min_images": None,
        "max_images": 0,
        "min_words": None,
        "max_words": 0,
        "dashboard_missing": 0,
        "historical_available": 0,
        "wikipedia_verified": 0,
    }

    expected_names = {row["file_name"] for row in rows}
    actual_names = {path.name for path in args.docx_dir.glob("*.docx")}
    web_names = {path.name for path in args.web_dir.glob("*.docx")}
    if expected_names != actual_names:
        errors.append("Output DOCX filenames do not match manifest")
    if expected_names != web_names:
        errors.append("Web DOCX filenames do not match manifest")

    for row in rows:
        source = args.docx_dir / row["file_name"]
        web = args.web_dir / row["file_name"]
        if not source.exists() or not web.exists():
            errors.append(f"Missing file: {row['file_name']}")
            continue
        if digest(source) != digest(web):
            errors.append(f"Web copy differs: {row['file_name']}")
        try:
            document = Document(source)
        except Exception as exc:
            errors.append(f"Cannot open {row['file_name']}: {exc}")
            continue
        text = "\n".join(
            [paragraph.text for paragraph in document.paragraphs]
            + [
                cell.text
                for table in document.tables
                for row_item in table.rows
                for cell in row_item.cells
            ]
        )
        headings = {
            paragraph.text.strip()
            for paragraph in document.paragraphs
            if paragraph.style.name.startswith("Heading")
        }
        missing_headings = REQUIRED_HEADINGS - headings
        if missing_headings:
            errors.append(
                f"{row['file_name']} missing headings: {sorted(missing_headings)}"
            )
        forbidden = sorted(value for value in FORBIDDEN_TEXT if value in text)
        if forbidden:
            errors.append(f"{row['file_name']} contains forbidden text: {forbidden}")
        if text.count("EJEMPLO DE FORMATO · NO APROBADO") != 2:
            errors.append(f"{row['file_name']} must contain exactly two example rows")
        if row.get("content_version") != "3.0-cdm-ready":
            errors.append(f"{row['file_name']} has wrong manifest content version")
        if "{{" in text or "}}" in text:
            errors.append(f"{row['file_name']} contains template tokens")
        if "para No disponible de los hogares" in text:
            errors.append(f"{row['file_name']} contains an invalid missing-value sentence")
        if row["dashboard_available"] and "No existe una serie municipal separada" in text:
            errors.append(f"{row['file_name']} incorrectly reports missing dashboard data")
        if not row["dashboard_available"] and "no se trasladaron cifras del municipio de origen" not in text:
            errors.append(f"{row['file_name']} missing safe source-gap warning")
        if row["municipio"] == "Villa Central":
            if row.get("historical_pmd_count") != 0:
                errors.append("Villa Central incorrectly counts the former district plan as a municipal PMD")
            if row.get("territorial_antecedent_count") != 1:
                errors.append("Villa Central predecessor-district antecedent is missing")
            if "Plan del antiguo Distrito Municipal de Villa Central" not in text:
                errors.append("Villa Central predecessor-district label is missing")

        paragraphs = len(document.paragraphs)
        tables = len(document.tables)
        images = len(document.inline_shapes)
        words = len(text.split())
        if images < 5:
            errors.append(f"{row['file_name']} has only {images} diagnostic images")
        if row["dashboard_available"] and words < 1800:
            errors.append(f"{row['file_name']} is too short for a completed diagnosis ({words} words)")
        if not row["dashboard_available"] and words < 900:
            errors.append(f"{row['file_name']} is too short for a source-gap draft ({words} words)")
        stats["min_paragraphs"] = (
            paragraphs
            if stats["min_paragraphs"] is None
            else min(stats["min_paragraphs"], paragraphs)
        )
        stats["max_paragraphs"] = max(stats["max_paragraphs"], paragraphs)
        stats["min_tables"] = (
            tables if stats["min_tables"] is None else min(stats["min_tables"], tables)
        )
        stats["max_tables"] = max(stats["max_tables"], tables)
        stats["min_images"] = images if stats["min_images"] is None else min(stats["min_images"], images)
        stats["max_images"] = max(stats["max_images"], images)
        stats["min_words"] = words if stats["min_words"] is None else min(stats["min_words"], words)
        stats["max_words"] = max(stats["max_words"], words)
        stats["dashboard_missing"] += int(not row["dashboard_available"])
        stats["historical_available"] += int(row["historical_pmd_count"] > 0)
        stats["wikipedia_verified"] += int(row["wikipedia_status"] == "verified")

    result = {
        "ok": not errors,
        "stats": stats,
        "errors": errors[:30],
        "error_count": len(errors),
    }
    print(json.dumps(result, ensure_ascii=False, indent=2))
    return 0 if not errors else 1


if __name__ == "__main__":
    raise SystemExit(main())
