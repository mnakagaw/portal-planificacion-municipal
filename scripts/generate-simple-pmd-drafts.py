#!/usr/bin/env python3
"""Generate concise, source-traceable municipal PMD draft DOCX files.

The generator deliberately avoids unrestricted web sources. It uses:
  * the municipal dashboard datasets (official statistics);
  * locally inventoried historical PMDs;
  * a short, verified Spanish Wikipedia lead for basic geographic context.

It never invents meetings, participation, approvals, budgets, beneficiaries,
authorities, or completed projects. Strategic content is explicitly labelled
as a proposal for OMPP/CDM review.
"""

from __future__ import annotations

import argparse
import csv
import json
import math
import re
import shutil
import sys
import time
import unicodedata
import urllib.parse
import urllib.request
from collections import defaultdict
from concurrent.futures import ThreadPoolExecutor, as_completed
from dataclasses import dataclass
from datetime import date
from pathlib import Path
from typing import Any, Iterable

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_ALIGN_VERTICAL, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


PERIOD = "2025-2028"
TODAY = date(2026, 7, 27)
USER_AGENT = "DDPT-PMD-Draft-Builder/1.0 (municipal planning research)"

COLORS = {
    "ink": "17352F",
    "teal": "347C6C",
    "teal_dark": "245B50",
    "slate": "465A55",
    "muted": "66736F",
    "line": "CCD8D4",
    "soft": "EDF3F1",
    "soft_blue": "EEF3F8",
    "warn": "8A5A20",
    "warn_soft": "FBF3E6",
    "white": "FFFFFF",
}

REGION_CODES = {
    "Cibao Norte": "01",
    "Cibao Sur": "02",
    "Cibao Nordeste": "03",
    "Cibao Noroeste": "04",
    "Valdesia": "05",
    "Enriquillo": "06",
    "El Valle": "07",
    "Yuma": "08",
    "Higuamo": "09",
    "Ozama": "10",
}

ALIASES = {
    "santodomingodeguzman": "distritonacional",
    "azuadecompostela": "azua",
    "concepciondelavega": "lavega",
    "hatomayordelrey": "hatomayor",
    "lamata": "villalamata",
    "neiba": "neyba",
    "sanfelipedepuertoplata": "puertoplata",
    "sanjosedelosllanos": "losllanos",
    "sanjuandelamaguana": "sanjuan",
    "santabarbaradesamana": "samana",
    "santiagodeloscaballeros": "santiago",
    "villabisononavarrete": "bisono",
}


def clean(value: str | None) -> str:
    text = unicodedata.normalize("NFD", value or "")
    text = "".join(ch for ch in text if unicodedata.category(ch) != "Mn")
    return re.sub(r"[^a-z0-9]", "", text.lower())


def canonical_municipality(value: str | None) -> str:
    key = clean(value)
    return ALIASES.get(key, key)


def canonical_province(value: str | None) -> str:
    key = clean(value)
    return "baoruco" if key == "bahoruco" else key


def territory_key(municipality: str, province: str) -> str:
    return f"{canonical_municipality(municipality)}|{canonical_province(province)}"


def code_key(value: Any) -> str:
    if value is None:
        return ""
    try:
        return f"{int(value):05d}"
    except (TypeError, ValueError):
        return str(value).zfill(5)


def slugify(value: str) -> str:
    text = unicodedata.normalize("NFD", value)
    text = "".join(ch for ch in text if unicodedata.category(ch) != "Mn")
    text = re.sub(r"[^A-Za-z0-9]+", "-", text).strip("-").lower()
    return text or "municipio"


def load_json(path: Path) -> Any:
    with path.open("r", encoding="utf-8") as handle:
        return json.load(handle)


def save_json(path: Path, value: Any) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    with path.open("w", encoding="utf-8", newline="\n") as handle:
        json.dump(value, handle, ensure_ascii=False, indent=2)
        handle.write("\n")


def percentage(part: float | int | None, total: float | int | None) -> float | None:
    if part is None or not total:
        return None
    return (float(part) / float(total)) * 100


def fmt_number(value: float | int | None, decimals: int = 0) -> str:
    if value is None:
        return "No disponible"
    if isinstance(value, float) and (math.isnan(value) or math.isinf(value)):
        return "No disponible"
    if decimals:
        return f"{float(value):,.{decimals}f}"
    return f"{float(value):,.0f}"


def fmt_pct(value: float | int | None, decimals: int = 1) -> str:
    if value is None:
        return "No disponible"
    return f"{float(value):.{decimals}f}%"


def first_sentence(text: str, max_chars: int = 430) -> str:
    compact = re.sub(r"\s+", " ", text or "").strip()
    compact = re.sub(r"\[[^\]]+\]", "", compact)
    if not compact:
        return ""
    pieces = re.split(r"(?<=[.!?])\s+", compact)
    result = pieces[0]
    if len(result) < 120 and len(pieces) > 1:
        result = f"{result} {pieces[1]}"
    return result[:max_chars].rstrip()


def fetch_wikipedia(
    municipality: str, province: str, cache: dict[str, Any]
) -> dict[str, Any]:
    key = territory_key(municipality, province)
    if key in cache and cache[key].get("status") in {"verified", "not_found"}:
        return cache[key]

    queries = [
        f'intitle:"{municipality}" "{province}" República Dominicana municipio',
        f'"{municipality}" "{province}" República Dominicana',
    ]
    result: dict[str, Any] = {
        "status": "not_found",
        "title": "",
        "url": "",
        "extract": "",
        "revid": None,
        "retrieved_at": TODAY.isoformat(),
    }
    for query in queries:
        params = {
            "action": "query",
            "generator": "search",
            "gsrsearch": query,
            "gsrlimit": "5",
            "prop": "extracts|info|revisions",
            "exintro": "1",
            "explaintext": "1",
            "inprop": "url",
            "rvprop": "ids",
            "format": "json",
            "formatversion": "2",
        }
        url = "https://es.wikipedia.org/w/api.php?" + urllib.parse.urlencode(params)
        try:
            request = urllib.request.Request(url, headers={"User-Agent": USER_AGENT})
            with urllib.request.urlopen(request, timeout=20) as response:
                payload = json.load(response)
            time.sleep(0.35)
        except Exception as exc:  # network failures are recorded, not hidden
            result["status"] = "error"
            result["error"] = str(exc)[:200]
            time.sleep(1.0)
            continue

        pages = payload.get("query", {}).get("pages", [])
        for page in pages:
            extract = page.get("extract", "") or ""
            normalized = clean(extract)
            if "republicadominicana" not in normalized:
                continue
            title_key = canonical_municipality(page.get("title", ""))
            municipality_key = canonical_municipality(municipality)
            province_key = canonical_province(province)
            if municipality_key not in title_key and municipality_key not in normalized:
                continue
            if province_key not in normalized and clean(province) not in clean(page.get("title", "")):
                continue
            result = {
                "status": "verified",
                "title": page.get("title", ""),
                "url": page.get("fullurl", ""),
                "extract": first_sentence(extract),
                "revid": (page.get("revisions") or [{}])[0].get("revid"),
                "retrieved_at": TODAY.isoformat(),
            }
            cache[key] = result
            return result

    cache[key] = result
    return result


def fetch_wikipedia_exact_batches(
    municipalities: list[dict[str, Any]], cache: dict[str, Any]
) -> None:
    """Resolve exact titles and redirects in small batches before text search."""
    pending = [
        item
        for item in municipalities
        if cache.get(territory_key(item["municipio"], item["provincia"]), {}).get("status")
        != "verified"
    ]
    for offset in range(0, len(pending), 35):
        batch = pending[offset : offset + 35]
        originals = [item["municipio"] for item in batch]
        params = {
            "action": "query",
            "titles": "|".join(originals),
            "redirects": "1",
            "prop": "extracts|info|revisions",
            "exintro": "1",
            "explaintext": "1",
            "inprop": "url",
            "rvprop": "ids",
            "format": "json",
            "formatversion": "2",
        }
        url = "https://es.wikipedia.org/w/api.php?" + urllib.parse.urlencode(params)
        try:
            request = urllib.request.Request(url, headers={"User-Agent": USER_AGENT})
            with urllib.request.urlopen(request, timeout=30) as response:
                payload = json.load(response)
            time.sleep(1.0)
        except Exception:
            time.sleep(2.0)
            continue

        query = payload.get("query", {})
        normalized = {item["from"]: item["to"] for item in query.get("normalized", [])}
        redirects = {item["from"]: item["to"] for item in query.get("redirects", [])}
        page_by_title = {page.get("title", ""): page for page in query.get("pages", [])}
        for item in batch:
            original = item["municipio"]
            resolved = normalized.get(original, original)
            resolved = redirects.get(resolved, resolved)
            page = page_by_title.get(resolved)
            if not page or page.get("missing"):
                continue
            extract = page.get("extract", "") or ""
            normalized_extract = clean(extract)
            municipality_key = canonical_municipality(original)
            province_key = canonical_province(item["provincia"])
            if "republicadominicana" not in normalized_extract:
                continue
            if municipality_key not in canonical_municipality(page.get("title", "")) and municipality_key not in normalized_extract:
                continue
            if province_key not in normalized_extract and clean(item["provincia"]) not in clean(page.get("title", "")):
                continue
            cache[territory_key(original, item["provincia"])] = {
                "status": "verified",
                "title": page.get("title", ""),
                "url": page.get("fullurl", ""),
                "extract": first_sentence(extract),
                "revid": (page.get("revisions") or [{}])[0].get("revid"),
                "retrieved_at": TODAY.isoformat(),
            }


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=90, start=110, bottom=90, end=110) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color: str = "CCD8D4", size: int = 5) -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = borders.find(qn(f"w:{edge}"))
        if tag is None:
            tag = OxmlElement(f"w:{edge}")
            borders.append(tag)
        tag.set(qn("w:val"), "single")
        tag.set(qn("w:sz"), str(size))
        tag.set(qn("w:color"), color)


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def add_text(
    paragraph,
    text: str,
    *,
    bold: bool = False,
    color: str | None = None,
    size: float | None = None,
    italic: bool = False,
) -> None:
    run = paragraph.add_run(text)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    if size:
        run.font.size = Pt(size)


def add_source_note(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="Source Note")
    add_text(p, f"Fuente: {text}", color=COLORS["muted"], size=8.5)


def add_callout(doc: Document, title: str, body: str, tone: str = "info") -> None:
    table = doc.add_table(rows=1, cols=1)
    table.autofit = True
    cell = table.cell(0, 0)
    set_cell_shading(cell, COLORS["warn_soft"] if tone == "warn" else COLORS["soft"])
    set_cell_margins(cell, top=150, start=180, bottom=150, end=180)
    p = cell.paragraphs[0]
    add_text(
        p,
        title,
        bold=True,
        color=COLORS["warn"] if tone == "warn" else COLORS["teal_dark"],
        size=10,
    )
    p = cell.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    add_text(p, body, color=COLORS["ink"], size=9.5)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_data_table(
    doc: Document,
    headers: list[str],
    rows: Iterable[Iterable[Any]],
    widths_cm: list[float] | None = None,
) -> None:
    rows = [list(row) for row in rows]
    table = doc.add_table(rows=1, cols=len(headers))
    table.autofit = False
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_table_borders(table)
    header = table.rows[0]
    set_repeat_table_header(header)
    for i, text in enumerate(headers):
        cell = header.cells[i]
        set_cell_shading(cell, COLORS["ink"])
        set_cell_margins(cell)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        add_text(p, str(text), bold=True, color=COLORS["white"], size=8.6)
        if widths_cm:
            cell.width = Cm(widths_cm[i])
    for row_index, values in enumerate(rows):
        cells = table.add_row().cells
        for i, value in enumerate(values):
            cell = cells[i]
            if row_index % 2:
                set_cell_shading(cell, "F7F9F8")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            p = cell.paragraphs[0]
            add_text(p, str(value), color=COLORS["ink"], size=8.5)
            if widths_cm:
                cell.width = Cm(widths_cm[i])
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def configure_document(doc: Document, municipality: str, province: str, code: str) -> None:
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(1.55)
    section.bottom_margin = Cm(1.55)
    section.left_margin = Cm(1.65)
    section.right_margin = Cm(1.65)
    section.header_distance = Cm(0.65)
    section.footer_distance = Cm(0.65)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Aptos"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Aptos")
    normal.font.size = Pt(9.5)
    normal.font.color.rgb = RGBColor.from_string(COLORS["ink"])
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.line_spacing = 1.05

    for style_name, size, color, before, after in (
        ("Title", 28, COLORS["ink"], 0, 8),
        ("Subtitle", 14, COLORS["teal"], 0, 7),
        ("Heading 1", 17, COLORS["ink"], 4, 8),
        ("Heading 2", 11.5, COLORS["teal_dark"], 8, 4),
        ("Heading 3", 10, COLORS["slate"], 6, 3),
    ):
        style = styles[style_name]
        style.font.name = "Aptos Display" if style_name in {"Title", "Heading 1"} else "Aptos"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), style.font.name)
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = style_name != "Subtitle"
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    if "Source Note" not in styles:
        style = styles.add_style("Source Note", WD_STYLE_TYPE.PARAGRAPH)
        style.font.name = "Aptos"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Aptos")
        style.font.size = Pt(8.5)
        style.font.color.rgb = RGBColor.from_string(COLORS["muted"])
        style.paragraph_format.space_before = Pt(2)
        style.paragraph_format.space_after = Pt(7)

    header = section.header
    table = header.add_table(rows=1, cols=2, width=Cm(17.7))
    table.autofit = False
    table.columns[0].width = Cm(12.8)
    table.columns[1].width = Cm(4.9)
    left = table.cell(0, 0).paragraphs[0]
    add_text(left, "PLANIFICACIÓN MUNICIPAL  |  BORRADOR TÉCNICO", bold=True, color=COLORS["teal_dark"], size=7.5)
    right = table.cell(0, 1).paragraphs[0]
    right.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    add_text(right, f"{municipality}  ·  {code or 'código por confirmar'}", color=COLORS["muted"], size=7.5)
    border = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), COLORS["line"])
    border.append(bottom)
    table.cell(0, 0).paragraphs[0]._p.get_or_add_pPr().append(border)

    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    add_text(p, f"{municipality}, {province}  |  PMD {PERIOD}  |  ", color=COLORS["muted"], size=7.5)
    add_text(p, "Borrador no aprobado", bold=True, color=COLORS["warn"], size=7.5)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_text(p, "    ·    ")
    field = OxmlElement("w:fldSimple")
    field.set(qn("w:instr"), "PAGE")
    p._p.append(field)


def page_break(doc: Document) -> None:
    p = doc.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)


@dataclass
class SourceBundle:
    basic: dict[str, Any] | None
    households: dict[str, Any] | None
    urban_rural: dict[str, Any] | None
    living: dict[str, Any] | None
    education: dict[str, Any] | None
    economy: dict[str, Any] | None
    tic: dict[str, Any] | None
    health: dict[str, Any] | None


def service_share(living: dict[str, Any] | None, service: str, category: str) -> float | None:
    if not living:
        return None
    item = living.get("servicios", {}).get(service, {})
    return percentage(item.get("categorias", {}).get(category), item.get("total"))


def make_findings(bundle: SourceBundle) -> list[dict[str, str]]:
    findings: list[dict[str, str]] = []
    basic = bundle.basic or {}
    variation = basic.get("variacion_pct")
    if variation is not None:
        direction = "aumentó" if variation >= 0 else "disminuyó"
        findings.append(
            {
                "title": "Cambio demográfico",
                "text": f"La población {direction} {abs(float(variation)):.1f}% entre 2010 y 2022. "
                "La programación de servicios debe revisar dónde se concentra ese cambio.",
                "evidence": "[DASH-BASE]",
            }
        )

    indoor_water = service_share(
        bundle.living, "agua_uso_domestico", "del_acueducto_dentro_de_la_vivienda"
    )
    if indoor_water is not None and indoor_water < 80:
        findings.append(
            {
                "title": "Acceso domiciliario al agua",
                "text": f"{indoor_water:.1f}% de los hogares reporta agua del acueducto dentro de la vivienda. "
                "La OMPP debe confirmar continuidad, calidad y distribución territorial del servicio.",
                "evidence": "[DASH-VIDA]",
            }
        )

    no_toilet = service_share(bundle.living, "servicios_sanitarios", "no_tiene")
    if no_toilet is not None and no_toilet >= 3:
        findings.append(
            {
                "title": "Saneamiento",
                "text": f"{no_toilet:.1f}% de los hogares no reporta servicio sanitario. "
                "El dato requiere localización territorial antes de definir acciones.",
                "evidence": "[DASH-VIDA]",
            }
        )

    internet_rate = ((bundle.tic or {}).get("internet") or {}).get("rate_used")
    computer_rate = ((bundle.tic or {}).get("computer") or {}).get("rate_used")
    if internet_rate is not None and float(internet_rate) < 0.70:
        computer_text = (
            f"; el uso de computadora alcanza {float(computer_rate) * 100:.1f}%"
            if computer_rate is not None
            else ""
        )
        findings.append(
            {
                "title": "Inclusión digital",
                "text": f"El uso de internet alcanza {float(internet_rate) * 100:.1f}%{computer_text}. "
                "La brecha debe analizarse por edad, sexo y zona cuando exista información.",
                "evidence": "[DASH-TIC]",
            }
        )

    secondary = (
        (((bundle.education or {}).get("anuario") or {}).get("eficiencia") or {}).get("secundario")
        or {}
    )
    dropout = secondary.get("abandono")
    if dropout is not None and float(dropout) >= 3:
        findings.append(
            {
                "title": "Permanencia escolar",
                "text": f"El abandono reportado en secundaria es {float(dropout):.1f}% para el distrito educativo asociado. "
                "El valor debe interpretarse como referencia distrital y no como resultado exclusivo del municipio.",
                "evidence": "[DASH-EDU]",
            }
        )

    return findings[:5]


def project_ideas(bundle: SourceBundle, findings: list[dict[str, str]]) -> list[list[str]]:
    ideas: list[list[str]] = []
    titles = {item["title"] for item in findings}
    if "Acceso domiciliario al agua" in titles:
        ideas.append(
            [
                "Diagnóstico operativo de agua y saneamiento",
                "Mapear continuidad, calidad, cobertura y puntos críticos antes de formular obras.",
                "[DASH-VIDA]",
                "Coordinación municipal-sectorial",
            ]
        )
    if "Saneamiento" in titles:
        ideas.append(
            [
                "Plan focalizado de saneamiento domiciliario",
                "Identificar hogares y sectores sin servicio sanitario y acordar una ruta institucional.",
                "[DASH-VIDA]",
                "Coordinación municipal-sectorial",
            ]
        )
    waste_collection = service_share(
        bundle.living, "eliminacion_basura", "la_recoge_el_ayuntamiento"
    )
    if waste_collection is not None and waste_collection < 90:
        ideas.append(
            [
                "Mejora de cobertura de residuos sólidos",
                "Revisar rutas, frecuencia y sectores no cubiertos con registro verificable.",
                "[DASH-VIDA]",
                "Competencia municipal",
            ]
        )
    if "Inclusión digital" in titles:
        ideas.append(
            [
                "Puntos municipales de acceso y alfabetización digital",
                "Definir espacios y grupos prioritarios después de validar la brecha.",
                "[DASH-TIC]",
                "Municipal / alianza educativa",
            ]
        )
    if "Permanencia escolar" in titles:
        ideas.append(
            [
                "Mesa local de permanencia escolar",
                "Coordinar seguimiento con el distrito educativo, sin atribuir al ayuntamiento competencias escolares directas.",
                "[DASH-EDU]",
                "Coordinación interinstitucional",
            ]
        )
    ideas.append(
        [
            "Tablero municipal de indicadores del PMD",
            "Mantener una línea base anual con fuente, año, responsable y estado de verificación.",
            "[DASH-BASE]",
            "Competencia municipal",
        ]
    )
    return ideas[:5]


def build_document(
    municipality: dict[str, Any],
    geographic_code: str,
    bundle: SourceBundle,
    historical: list[dict[str, Any]],
    wikipedia: dict[str, Any],
    output_path: Path,
) -> None:
    name = municipality["municipio"]
    province = municipality["provincia"]
    region = municipality["region"]
    doc = Document()
    configure_document(doc, name, province, geographic_code)

    # Cover
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(32)
    add_text(p, "REPÚBLICA DOMINICANA", bold=True, color=COLORS["teal"], size=9)
    p = doc.add_paragraph(style="Title")
    add_text(p, "Plan Municipal\nde Desarrollo", bold=True, color=COLORS["ink"], size=30)
    p = doc.add_paragraph(style="Subtitle")
    add_text(p, name, bold=True, color=COLORS["teal_dark"], size=18)
    p = doc.add_paragraph()
    add_text(p, province, color=COLORS["muted"], size=11)
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(24)
    add_text(p, f"BORRADOR TÉCNICO  ·  {PERIOD}", bold=True, color=COLORS["white"], size=11)
    set_cell = doc.add_table(rows=1, cols=1).cell(0, 0)
    set_cell_shading(set_cell, COLORS["ink"])
    set_cell_margins(set_cell, top=180, start=200, bottom=180, end=200)
    cover_p = set_cell.paragraphs[0]
    add_text(cover_p, f"Borrador técnico  |  Período {PERIOD}", bold=True, color=COLORS["white"], size=12)
    cover_p = set_cell.add_paragraph()
    add_text(
        cover_p,
        "Documento de apoyo para revisión de la OMPP y el CDM. No constituye un PMD aprobado.",
        color=COLORS["white"],
        size=9.5,
    )
    doc.add_paragraph()
    add_callout(
        doc,
        "Alcance del documento",
        "Se elaboró únicamente con datos del Dashboard territorial, PMD anteriores disponibles y "
        "Wikipedia como referencia secundaria de ubicación. No acredita participación ciudadana, "
        "reuniones, decisiones del CDM ni aprobación del Concejo de Regidores.",
        "warn",
    )
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(22)
    add_text(p, f"Código geográfico: {geographic_code or 'por confirmar'}", color=COLORS["muted"], size=9)
    add_text(p, f"\nFecha de corte de fuentes: {TODAY.strftime('%d-%m-%Y')}", color=COLORS["muted"], size=9)

    page_break(doc)
    doc.add_heading("Cómo usar este borrador", level=1)
    add_callout(
        doc,
        "Regla principal",
        "Conservar únicamente frases que puedan rastrearse a una fuente. Toda corrección municipal "
        "debe registrar responsable, fecha y evidencia.",
    )
    bullets = [
        ("Información general y diagnóstico", "Revisar cifras, años, unidades y cobertura territorial."),
        ("Hallazgos y FODA", "Confirmar, corregir o eliminar. No representan todavía opinión ciudadana."),
        ("Visión y líneas estratégicas", "Son una base de discusión; deben redactarse mediante el proceso municipal."),
        ("Ideas de proyectos", "No contienen presupuesto, aprobación ni compromiso institucional."),
    ]
    for title, body in bullets:
        p = doc.add_paragraph(style="List Bullet")
        add_text(p, f"{title}: ", bold=True, color=COLORS["teal_dark"])
        add_text(p, body)
    doc.add_heading("Ficha de control", level=2)
    add_data_table(
        doc,
        ["Campo", "Valor"],
        [
            ["Estado", "Borrador técnico no aprobado"],
            ["Municipio / provincia", f"{name} / {province}"],
            ["Región", region],
            ["Período propuesto", PERIOD],
            ["Código geográfico", geographic_code or "Por confirmar"],
            ["Revisión requerida", "OMPP y CDM"],
        ],
        [5.1, 12.2],
    )
    doc.add_heading("Contenido", level=2)
    for line in [
        "1. Información general",
        "2. Diagnóstico municipal",
        "3. Hallazgos prioritarios",
        "4. FODA preliminar",
        "5. Visión y líneas estratégicas",
        "6. Ideas de proyectos",
        "7. Revisión y fuentes",
    ]:
        doc.add_paragraph(line, style="List Number")

    page_break(doc)
    doc.add_heading("1. Información general", level=1)
    if wikipedia.get("status") == "verified" and wikipedia.get("extract"):
        doc.add_paragraph(f"{wikipedia['extract']} [WIKI-01]")
        add_source_note(
            doc,
            f"Wikipedia en español, «{wikipedia.get('title')}», revisión {wikipedia.get('revid') or 'no disponible'}, "
            f"consulta {wikipedia.get('retrieved_at')}. Uso limitado a contexto geográfico básico.",
        )
    else:
        doc.add_paragraph(
            f"{name} se registra en el portal territorial como municipio de la provincia {province}, "
            f"región {region}. [PORTAL-01]"
        )
        add_source_note(doc, "Portal de Planificación Municipal, ficha territorial.")

    basic = bundle.basic or {}
    households = bundle.households or {}
    urban = bundle.urban_rural or {}
    population = basic.get("poblacion_total")
    population_2010 = basic.get("poblacion_2010")
    male_pct = percentage(basic.get("poblacion_hombres"), population)
    female_pct = percentage(basic.get("poblacion_mujeres"), population)
    urban_pct = percentage(urban.get("urbana"), urban.get("poblacion_total"))
    rural_pct = percentage(urban.get("rural"), urban.get("poblacion_total"))

    doc.add_heading("Indicadores de referencia", level=2)
    add_data_table(
        doc,
        ["Indicador", "Valor", "Año / fuente"],
        [
            ["Población total", fmt_number(population), "2022 · [DASH-BASE]"],
            ["Población 2010", fmt_number(population_2010), "2010 · [DASH-BASE]"],
            ["Variación 2010-2022", fmt_pct(basic.get("variacion_pct")), "[DASH-BASE]"],
            ["Hombres / mujeres", f"{fmt_pct(male_pct)} / {fmt_pct(female_pct)}", "2022 · [DASH-BASE]"],
            ["Población urbana / rural", f"{fmt_pct(urban_pct)} / {fmt_pct(rural_pct)}", "2022 · [DASH-URB]"],
            ["Hogares", fmt_number(households.get("hogares_total")), "2022 · [DASH-HOG]"],
            ["Personas por hogar", fmt_number(households.get("personas_por_hogar"), 2), "2022 · [DASH-HOG]"],
            ["Viviendas ocupadas", fmt_number(basic.get("viviendas_ocupadas")), "2022 · [DASH-BASE]"],
        ],
        [7.2, 4.5, 5.6],
    )
    add_callout(
        doc,
        "Lectura breve",
        (
            f"El municipio registra {fmt_number(population)} habitantes en 2022. "
            f"La variación respecto de 2010 es {fmt_pct(basic.get('variacion_pct'))}. "
            "Estos valores son una línea base estadística; no explican por sí solos las causas del cambio."
        )
        if population is not None
        else "El Dashboard no ofrece una ficha municipal completa. La OMPP debe completar la línea base con una fuente oficial.",
    )

    doc.add_heading("Antecedentes de planificación", level=2)
    if historical:
        add_data_table(
            doc,
            ["Documento localizado", "Período", "Condición", "Uso en este borrador"],
            [
                [
                    f"PMD de {name}",
                    f"{row.get('period_start')}-{row.get('period_end')}",
                    "Oficial" if row.get("document_kind") == "official" else "Borrador",
                    "Antecedente para revisar continuidad; no se presume vigente",
                ]
                for row in sorted(historical, key=lambda item: item.get("period_end") or 0, reverse=True)
            ],
            [5.1, 2.6, 2.5, 7.1],
        )
        add_source_note(
            doc,
            "Inventario local de PMD históricos. Los proyectos anteriores deben clasificarse como "
            "concluidos, en curso, pendientes o descartados antes de trasladarlos al nuevo plan.",
        )
    else:
        doc.add_paragraph(
            "No se localizó un PMD anterior en el inventario disponible. Esto no prueba que no exista; "
            "la OMPP debe verificar su archivo institucional."
        )

    page_break(doc)
    doc.add_heading("2. Diagnóstico municipal", level=1)
    doc.add_paragraph(
        "El diagnóstico resume indicadores verificables y evita interpretar como hechos los asuntos que "
        "requieren trabajo de campo o validación municipal."
    )

    doc.add_heading("2.1 Vivienda y servicios", level=2)
    indoor_water = service_share(
        bundle.living, "agua_uso_domestico", "del_acueducto_dentro_de_la_vivienda"
    )
    toilet = service_share(bundle.living, "servicios_sanitarios", "inodoro")
    no_toilet = service_share(bundle.living, "servicios_sanitarios", "no_tiene")
    grid = service_share(bundle.living, "alumbrado", "energia_electrica_del_tendido_publico")
    waste = service_share(bundle.living, "eliminacion_basura", "la_recoge_el_ayuntamiento")
    bottled = service_share(bundle.living, "agua_para_beber", "botellones")
    add_data_table(
        doc,
        ["Indicador", "Valor", "Lectura prudente"],
        [
            ["Agua del acueducto dentro de la vivienda", fmt_pct(indoor_water), "Cobertura declarada por hogares"],
            ["Hogares con inodoro", fmt_pct(toilet), "No informa calidad del sistema"],
            ["Hogares sin servicio sanitario", fmt_pct(no_toilet), "Requiere localización territorial"],
            ["Electricidad de red pública", fmt_pct(grid), "Fuente principal declarada"],
            ["Basura recogida por el ayuntamiento", fmt_pct(waste), "No informa frecuencia ni calidad"],
            ["Agua de beber en botellones", fmt_pct(bottled), "Patrón de abastecimiento declarado"],
        ],
        [7.0, 3.1, 7.2],
    )
    add_source_note(doc, "Dashboard territorial, Censo Nacional de Población y Vivienda 2022. [DASH-VIDA]")

    doc.add_heading("2.2 Educación, salud y conectividad", level=2)
    education = bundle.education or {}
    efficiency = ((education.get("anuario") or {}).get("eficiencia") or {})
    secondary = efficiency.get("secundario") or {}
    primary = efficiency.get("primario") or {}
    internet = ((bundle.tic or {}).get("internet") or {}).get("rate_used")
    computer = ((bundle.tic or {}).get("computer") or {}).get("rate_used")
    health = bundle.health or {}
    health_count = (
        health.get("total_establecimientos")
        or health.get("total")
        or health.get("count")
        or (len(health.get("centros", [])) if health.get("centros") is not None else None)
    )
    add_data_table(
        doc,
        ["Indicador", "Valor", "Cobertura"],
        [
            ["Promoción primaria", fmt_pct(primary.get("promocion")), "Distrito educativo asociado · [DASH-EDU]"],
            ["Abandono secundario", fmt_pct(secondary.get("abandono")), "Distrito educativo asociado · [DASH-EDU]"],
            ["Uso de internet", fmt_pct(float(internet) * 100 if internet is not None else None), "Población de referencia · [DASH-TIC]"],
            ["Uso de computadora", fmt_pct(float(computer) * 100 if computer is not None else None), "Población de referencia · [DASH-TIC]"],
            ["Establecimientos de salud registrados", fmt_number(health_count), "Registro del Dashboard · [DASH-SALUD]"],
        ],
        [7.0, 3.3, 7.0],
    )
    add_callout(
        doc,
        "Precaución",
        "Los indicadores educativos pueden corresponder al distrito educativo y no exclusivamente al municipio. "
        "Los establecimientos de salud no equivalen a capacidad, calidad ni cobertura efectiva.",
        "warn",
    )

    page_break(doc)
    doc.add_heading("2.3 Economía y empleo", level=2)
    economy = bundle.economy or {}
    dee = economy.get("dee_2024") or {}
    top = dee.get("top_specialization") or {}
    add_data_table(
        doc,
        ["Indicador", "Valor", "Año / fuente"],
        [
            ["Establecimientos económicos", fmt_number(dee.get("total_establishments")), "DEE 2024 · [DASH-ECO]"],
            ["Empleo estimado", fmt_number(dee.get("total_employees"), 1), "DEE 2024 · [DASH-ECO]"],
            ["Empleo medio por establecimiento", fmt_number(dee.get("avg_employees_per_establishment"), 1), "DEE 2024 · [DASH-ECO]"],
            ["Especialización relativa principal", top.get("label") or "No disponible", "DEE 2024 · [DASH-ECO]"],
        ],
        [7.0, 5.0, 5.3],
    )
    doc.add_paragraph(
        "El directorio empresarial describe establecimientos formales observados por la fuente. No debe "
        "interpretarse como el total de empleo ni de actividad económica del municipio."
    )

    doc.add_heading("3. Hallazgos prioritarios", level=1)
    findings = make_findings(bundle)
    if findings:
        for index, item in enumerate(findings, 1):
            doc.add_heading(f"{index}. {item['title']}", level=2)
            doc.add_paragraph(f"{item['text']} {item['evidence']}")
    else:
        add_callout(
            doc,
            "Línea base incompleta",
            "No hay datos suficientes para proponer hallazgos automáticos. La OMPP debe completar la ficha "
            "con fuentes oficiales antes de formular prioridades.",
            "warn",
        )
    doc.add_heading("Continuidad con PMD anteriores", level=2)
    if historical:
        doc.add_paragraph(
            "Los PMD anteriores localizados deben revisarse para distinguir proyectos concluidos, en ejecución, "
            "pendientes y descartados. Este borrador no traslada automáticamente proyectos ni compromisos de "
            f"los períodos {', '.join(f'{r.get('period_start')}-{r.get('period_end')}' for r in historical)}. "
            "[PMD-ANT]"
        )
    else:
        doc.add_paragraph(
            "No se incorporan supuestos sobre planes anteriores. La OMPP debe documentar si existe un antecedente "
            "no incluido en el inventario."
        )

    page_break(doc)
    doc.add_heading("4. FODA preliminar", level=1)
    doc.add_paragraph(
        "Esta matriz es una ayuda técnica. Fortalezas y debilidades se apoyan en indicadores; oportunidades y "
        "amenazas quedan como preguntas porque requieren evidencia externa o deliberación local."
    )
    strengths: list[str] = []
    weaknesses: list[str] = []
    if grid is not None and grid >= 90:
        strengths.append(f"Alta cobertura declarada de electricidad de red ({grid:.1f}%). [DASH-VIDA]")
    if waste is not None and waste >= 85:
        strengths.append(f"Cobertura municipal declarada de recogida de residuos ({waste:.1f}%). [DASH-VIDA]")
    if urban_pct is not None:
        strengths.append(f"Perfil territorial cuantificado: {urban_pct:.1f}% urbano y {rural_pct:.1f}% rural. [DASH-URB]")
    for item in findings:
        if item["title"] in {"Acceso domiciliario al agua", "Saneamiento", "Inclusión digital", "Permanencia escolar"}:
            weaknesses.append(f"{item['title']}: {item['text']} {item['evidence']}")
    if not strengths:
        strengths.append("Por confirmar con evidencia municipal.")
    if not weaknesses:
        weaknesses.append("Por confirmar después de revisar la línea base.")
    add_data_table(
        doc,
        ["Fortalezas verificables", "Debilidades verificables"],
        [["\n".join(f"• {x}" for x in strengths), "\n".join(f"• {x}" for x in weaknesses)]],
        [8.65, 8.65],
    )
    add_data_table(
        doc,
        ["Oportunidades por validar", "Amenazas por validar"],
        [
            [
                "¿Qué programas públicos, alianzas o inversiones vigentes pueden responder a los hallazgos?",
                "¿Qué riesgos ambientales, fiscales, institucionales o económicos cuentan con evidencia actual?",
            ]
        ],
        [8.65, 8.65],
    )

    page_break(doc)
    doc.add_heading("5. Visión y líneas estratégicas", level=1)
    add_callout(
        doc,
        "Propuesta base para discusión",
        f"Al 2028, {name} será un municipio que mejora de forma medible sus brechas prioritarias, "
        "fortalece los activos que la comunidad valide y gestiona su territorio con transparencia, "
        "inclusión y seguimiento de resultados.",
        "warn",
    )
    doc.add_paragraph(
        "La frase anterior no es una visión aprobada. Debe reescribirse después de validar el diagnóstico y "
        "documentar el proceso participativo."
    )
    add_data_table(
        doc,
        ["Línea preliminar", "Propósito de revisión"],
        [
            ["1. Gestión municipal y datos", "Organizar responsabilidades, fuentes, metas y seguimiento."],
            ["2. Servicios y desarrollo social", "Priorizar brechas verificadas de vivienda, educación, salud y conectividad."],
            ["3. Economía local", "Precisar activos productivos, empleo y alianzas con evidencia."],
            ["4. Territorio y ambiente", "Identificar riesgos, uso del suelo y sostenibilidad mediante fuentes técnicas."],
        ],
        [6.2, 11.1],
    )
    doc.add_heading("Preguntas mínimas", level=2)
    for question in [
        "¿Qué dos cambios deben ser visibles y medibles al cierre de 2028?",
        "¿Qué activos del municipio están confirmados por evidencia y por la comunidad?",
        "¿Qué asuntos son competencia municipal y cuáles requieren gestión ante otra institución?",
    ]:
        doc.add_paragraph(question, style="List Bullet")

    page_break(doc)
    doc.add_heading("6. Ideas de proyectos", level=1)
    doc.add_paragraph(
        "Las siguientes ideas se derivan de brechas estadísticas. No son proyectos aprobados y no incluyen "
        "presupuesto, beneficiarios, cronograma ni fuente de financiamiento."
    )
    ideas = project_ideas(bundle, findings)
    add_data_table(
        doc,
        ["Idea preliminar", "Qué debe verificarse", "Evidencia", "Competencia inicial"],
        ideas,
        [4.7, 7.0, 2.1, 3.5],
    )
    doc.add_heading("Ficha para completar por proyecto", level=2)
    add_data_table(
        doc,
        ["Campo", "Registro municipal"],
        [
            ["Problema y evidencia", ""],
            ["Objetivo / resultado", ""],
            ["Responsable y aliados", ""],
            ["Competencia legal", ""],
            ["Meta, indicador y línea base", ""],
            ["Costo y fuente de financiamiento", ""],
            ["Estado de validación", ""],
        ],
        [6.1, 11.2],
    )

    page_break(doc)
    doc.add_heading("7. Revisión OMPP y CDM", level=1)
    add_data_table(
        doc,
        ["Paso", "Responsable", "Resultado esperado", "Fecha / evidencia"],
        [
            ["1. Confirmar cifras y fuentes", "OMPP", "Correcciones registradas", ""],
            ["2. Revisar hallazgos y FODA", "OMPP + áreas técnicas", "Hallazgos aceptados o eliminados", ""],
            ["3. Priorizar problemas y activos", "CDM", "Matriz de priorización documentada", ""],
            ["4. Formular visión y objetivos", "CDM + Ayuntamiento", "Redacción acordada", ""],
            ["5. Completar cartera y demandas", "OMPP + CDM", "Responsables, metas y competencias", ""],
            ["6. Tramitar aprobación", "Autoridad competente", "Resolución y publicación, si corresponde", ""],
        ],
        [4.2, 3.5, 6.1, 3.5],
    )
    doc.add_heading("Registro de correcciones", level=2)
    add_data_table(
        doc,
        ["Sección / dato", "Corrección", "Fuente o evidencia", "Responsable / fecha"],
        [["", "", "", ""] for _ in range(5)],
        [4.0, 4.8, 4.8, 3.7],
    )
    add_callout(
        doc,
        "Importante",
        "No completar como hechos la participación ciudadana, las reuniones, la constitución del CDM, "
        "la dotación de la OMPP ni la aprobación del plan sin evidencia documental.",
        "warn",
    )

    page_break(doc)
    doc.add_heading("Fuentes y trazabilidad", level=1)
    source_rows = [
        ["PORTAL-01", "Portal de Planificación Municipal", TODAY.isoformat(), "Ficha municipal y estado SISMAP", "Alta"],
        ["DASH-BASE", "Censo Nacional de Población y Vivienda", "2022", "Indicadores básicos", "Alta"],
        ["DASH-HOG", "Censo Nacional de Población y Vivienda", "2022", "Hogares", "Alta"],
        ["DASH-URB", "Censo Nacional de Población y Vivienda", "2022", "Población urbana y rural", "Alta"],
        ["DASH-VIDA", "Censo Nacional de Población y Vivienda", "2022", "Vivienda y servicios", "Alta"],
        ["DASH-EDU", "Anuario Estadístico Educativo", "2024", "Distrito educativo asociado", "Media"],
        ["DASH-TIC", "Censo Nacional de Población y Vivienda", "2022", "Uso de TIC", "Alta"],
        ["DASH-ECO", "Directorio de Empresas y Establecimientos", "2024", "Economía formal observada", "Media"],
        ["DASH-SALUD", "Registro presentado por el Dashboard", "s/f", "Establecimientos de salud", "Media"],
    ]
    if wikipedia.get("status") == "verified":
        source_rows.append(
            [
                "WIKI-01",
                wikipedia.get("url", "Wikipedia en español"),
                wikipedia.get("retrieved_at", TODAY.isoformat()),
                "Ubicación general; no se usa para cifras",
                "Baja",
            ]
        )
    for index, row in enumerate(
        sorted(historical, key=lambda item: item.get("period_end") or 0, reverse=True), 1
    ):
        page_ref = (
            f"Documento completo ({row.get('pages')} págs.)"
            if row.get("pages")
            else "Documento completo"
        )
        source_rows.append(
            [
                f"PMD-ANT-{index:02d}",
                row.get("source_url") or row.get("source_page") or row.get("local_path") or "PMD anterior",
                f"{row.get('period_start')}-{row.get('period_end')}",
                page_ref,
                "Alta" if row.get("validation_status", "").startswith("verified") else "Media",
            ]
        )
    add_data_table(
        doc,
        ["ID", "Fuente", "Año / consulta", "Uso o página", "Confianza"],
        source_rows,
        [2.0, 7.0, 2.7, 4.3, 1.5],
    )
    doc.add_heading("Reglas de redacción", level=2)
    for rule in [
        "Toda cifra debe conservar año, unidad y fuente.",
        "Una afirmación del PMD anterior debe escribirse en pasado y con período.",
        "Wikipedia solo se admite como referencia secundaria de contexto.",
        "Si una fuente no permite confirmar un dato, el dato se omite.",
        "Participación, acuerdos, aprobación y ejecución requieren evidencia municipal.",
    ]:
        doc.add_paragraph(rule, style="List Bullet")

    doc.core_properties.title = f"PMD de {name} - Borrador técnico {PERIOD}"
    doc.core_properties.subject = "Borrador municipal basado en fuentes verificables"
    doc.core_properties.author = "DDPT - Generación asistida para revisión municipal"
    doc.core_properties.keywords = "PMD, borrador, municipio, OMPP, CDM, trazabilidad"
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)


def read_classifier_codes(path: Path) -> dict[str, list[str]]:
    result: dict[str, list[str]] = defaultdict(list)
    if not path.exists():
        return result
    for raw in path.read_text(encoding="utf-8", errors="ignore").splitlines():
        match = re.match(r"^\s*(\d{8,10})\s*(.+?)\s*$", raw)
        if not match:
            continue
        result[canonical_municipality(match.group(2))].append(match.group(1))
    return result


def build_dataset_index(items: Any) -> dict[str, Any]:
    if isinstance(items, list):
        return {code_key(item.get("adm2_code")): item for item in items if item.get("adm2_code") is not None}
    if isinstance(items, dict):
        return {code_key(key): value for key, value in items.items()}
    return {}


def resolve_geographic_code(
    municipality: dict[str, Any], adm2_code: str, classifier: dict[str, list[str]]
) -> str:
    region_code = REGION_CODES.get(municipality.get("region", ""), "")
    if region_code and adm2_code:
        # Dashboard ADM2 uses PPMMM (province + three-digit municipality).
        # DIGEPRES uses RRPPMMMM (region + province + four-digit municipality).
        return f"{region_code}{adm2_code[:2]}0{adm2_code[2:]}"
    candidates = classifier.get(canonical_municipality(municipality.get("municipio", "")), [])
    return candidates[0] if len(candidates) == 1 else ""


def main() -> int:
    parser = argparse.ArgumentParser()
    parser.add_argument("--workspace", type=Path, required=True)
    parser.add_argument("--portal-repo", type=Path, required=True)
    parser.add_argument("--municipality", help="Generate one municipality (prototype mode)")
    parser.add_argument("--all", action="store_true", help="Generate all target municipalities")
    parser.add_argument("--skip-wikipedia", action="store_true")
    parser.add_argument("--web-copy", action="store_true", help="Copy DOCX files to portal public/downloads")
    parser.add_argument(
        "--resume",
        action="store_true",
        help="Reuse already generated DOCX files while rebuilding the complete manifest",
    )
    args = parser.parse_args()

    workspace = args.workspace.resolve()
    portal_repo = args.portal_repo.resolve()
    portal_data_path = portal_repo / "app" / "data" / "municipios.json"
    dashboard_dir = workspace / "tmp" / "dashboard_repo_full_inspect" / "src" / "data"
    inventory_path = workspace / "output" / "pmd_historico_fuentes_no_sismap" / "inventory.json"
    classifier_path = workspace / "output" / "clasificador_geografico" / "clasificador_geografico_01-10.txt"
    output_dir = workspace / "output" / "pmd_borradores_2025_2028"
    docx_dir = output_dir / "docx"
    cache_path = output_dir / "wikipedia_cache.json"
    manifest_path = output_dir / "manifest.json"
    web_dir = portal_repo / "public" / "downloads" / "pmd-borradores"

    portal_data = load_json(portal_data_path)
    targets = [
        item
        for item in portal_data
        if not item.get("pmd", {}).get("hasOfficialEvidence")
        and not (
            item.get("pmd", {}).get("has7_12")
            or item.get("pmd", {}).get("hasDraft")
        )
    ]
    if len(targets) != 104:
        raise RuntimeError(f"Expected 104 target municipalities; found {len(targets)}")

    if args.municipality:
        wanted = clean(args.municipality)
        targets = [item for item in targets if clean(item["municipio"]) == wanted]
        if not targets:
            raise RuntimeError(f"Target municipality not found: {args.municipality}")
    elif not args.all:
        parser.error("Choose --municipality NAME or --all")

    dashboard_index = load_json(dashboard_dir / "municipios_index.json")
    dashboard_by_territory = {
        territory_key(item["municipio"], item["provincia"]): code_key(item["adm2_code"])
        for item in dashboard_index
    }
    dataset_names = {
        "basic": "indicadores_basicos.json",
        "households": "hogares_resumen.json",
        "urban_rural": "poblacion_urbana_rural.json",
        "living": "condicion_vida.json",
        "education": "educacion.json",
        "economy": "economia_empleo.json",
        "tic": "tic.json",
        "health": "salud_establecimientos.json",
    }
    datasets = {
        name: build_dataset_index(load_json(dashboard_dir / filename))
        for name, filename in dataset_names.items()
    }

    historical_rows = load_json(inventory_path).get("rows", [])
    historical_by_territory: dict[str, list[dict[str, Any]]] = defaultdict(list)
    for row in historical_rows:
        if row.get("status") == "downloaded":
            historical_by_territory[territory_key(row["municipality"], row["province"])].append(row)

    classifier = read_classifier_codes(classifier_path)
    wikipedia_cache = load_json(cache_path) if cache_path.exists() else {}
    if not args.skip_wikipedia:
        fetch_wikipedia_exact_batches(targets, wikipedia_cache)
        save_json(cache_path, wikipedia_cache)
        unresolved = [
            item
            for item in targets
            if wikipedia_cache.get(
                territory_key(item["municipio"], item["provincia"]), {}
            ).get("status")
            != "verified"
        ]
        with ThreadPoolExecutor(max_workers=1) as executor:
            futures = {
                executor.submit(
                    fetch_wikipedia,
                    item["municipio"],
                    item["provincia"],
                    wikipedia_cache,
                ): item
                for item in unresolved
            }
            for future in as_completed(futures):
                future.result()
        save_json(cache_path, wikipedia_cache)

    manifest_rows: list[dict[str, Any]] = []
    docx_dir.mkdir(parents=True, exist_ok=True)
    if args.web_copy:
        web_dir.mkdir(parents=True, exist_ok=True)

    for item in targets:
        key = territory_key(item["municipio"], item["provincia"])
        adm2_code = dashboard_by_territory.get(key, "")
        geographic_code = resolve_geographic_code(item, adm2_code, classifier)
        bundle = SourceBundle(
            **{
                name: dataset.get(adm2_code) if adm2_code else None
                for name, dataset in datasets.items()
            }
        )
        historical = historical_by_territory.get(key, [])
        wikipedia = wikipedia_cache.get(key, {"status": "not_requested"})
        file_name = (
            f"{geographic_code or f'id-{item['id']:03d}'}_"
            f"PMD_{slugify(item['municipio'])}_Borrador_Tecnico_{PERIOD}.docx"
        )
        output_path = docx_dir / file_name
        if not (args.resume and output_path.exists()):
            build_document(item, geographic_code, bundle, historical, wikipedia, output_path)
        if args.web_copy:
            web_path = web_dir / file_name
            if not (args.resume and web_path.exists()):
                shutil.copy2(output_path, web_path)
        manifest_rows.append(
            {
                "id": item["id"],
                "municipio": item["municipio"],
                "provincia": item["provincia"],
                "region": item["region"],
                "geographic_code": geographic_code,
                "dashboard_adm2_code": adm2_code,
                "dashboard_available": bool(adm2_code),
                "historical_pmd_count": len(historical),
                "historical_periods": [
                    f"{row.get('period_start')}-{row.get('period_end')}" for row in historical
                ],
                "wikipedia_status": wikipedia.get("status"),
                "wikipedia_url": wikipedia.get("url", ""),
                "file_name": file_name,
                "relative_url": f"downloads/pmd-borradores/{file_name}",
            }
        )

    existing_manifest = []
    if manifest_path.exists() and args.municipality:
        existing_manifest = load_json(manifest_path).get("municipalities", [])
        generated_ids = {row["id"] for row in manifest_rows}
        existing_manifest = [row for row in existing_manifest if row["id"] not in generated_ids]
    combined = sorted(existing_manifest + manifest_rows, key=lambda row: row["id"])
    save_json(
        manifest_path,
        {
            "generated_at": TODAY.isoformat(),
            "period": PERIOD,
            "target_rule": "No PMD official evidence and no existing 7-12/draft",
            "expected_target_count": 104,
            "generated_count": len(combined),
            "municipalities": combined,
        },
    )
    csv_path = output_dir / "manifest.csv"
    csv_path.parent.mkdir(parents=True, exist_ok=True)
    with csv_path.open("w", encoding="utf-8-sig", newline="") as handle:
        columns = [
            "id",
            "municipio",
            "provincia",
            "region",
            "geographic_code",
            "dashboard_adm2_code",
            "dashboard_available",
            "historical_pmd_count",
            "historical_periods",
            "wikipedia_status",
            "wikipedia_url",
            "file_name",
            "relative_url",
        ]
        writer = csv.DictWriter(handle, fieldnames=columns)
        writer.writeheader()
        for row in combined:
            out = dict(row)
            out["historical_periods"] = "; ".join(out.get("historical_periods", []))
            writer.writerow(out)

    print(
        json.dumps(
            {
                "generated": len(manifest_rows),
                "manifest_total": len(combined),
                "output_dir": str(output_dir),
                "web_copy": args.web_copy,
            },
            ensure_ascii=False,
        )
    )
    return 0


if __name__ == "__main__":
    sys.exit(main())
