#!/usr/bin/env python3
"""Generate source-traceable PMD documents for final municipal review.

This is the third-generation builder for the Paquete Mínimo. It incorporates
the current PDF export of the territorial dashboard and adds municipal
narrative based only on:

* the dashboard's official statistical datasets;
* locally archived historical PMDs;
* Spanish Wikipedia as an explicitly secondary source.

Información General and Diagnóstico are written as near-final technical text.
Strengths and weaknesses are derived only from the diagnosis. Vision and the
action plan are left to the CDM, with one clearly marked example row. The script
does not invent meetings, participation, approvals, budgets, beneficiaries,
authorities or execution status.
"""

from __future__ import annotations

import argparse
import csv
import importlib.util
import json
import math
import re
import shutil
import subprocess
import sys
import textwrap
import time
import urllib.parse
import urllib.request
from collections import Counter, defaultdict
from datetime import date
from pathlib import Path
from typing import Any, Iterable

try:
    import fitz
except ImportError:  # Historical text stays available from the local cache.
    fitz = None
try:
    import matplotlib

    matplotlib.use("Agg")
    import matplotlib.pyplot as plt
    from matplotlib.patches import FancyBboxPatch, Polygon
except ImportError:
    matplotlib = None
    plt = None
    FancyBboxPatch = None
    Polygon = None

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


SCRIPT_DIR = Path(__file__).resolve().parent
BASE_GENERATOR_PATH = SCRIPT_DIR / "generate-simple-pmd-drafts.py"
spec = importlib.util.spec_from_file_location("simple_pmd_builder", BASE_GENERATOR_PATH)
if spec is None or spec.loader is None:
    raise RuntimeError(f"Could not load {BASE_GENERATOR_PATH}")
base = importlib.util.module_from_spec(spec)
sys.modules[spec.name] = base
spec.loader.exec_module(base)


PERIOD = "2025-2028"
TODAY = date.today()
CONTENT_VERSION = "3.2-dashboard-territorial"
USER_AGENT = "DDPT-PMD-Builder/3.2 (municipal planning research)"
ACTIVE_MAPA_STATES = {"EJECUCIÓN", "REPROGRAMAR"}
PROVINCE_ONLY_TERRITORIAL_KEYS = {
    "crime",
    "homicide",
    "traffic",
    "overcrowding",
    "disability",
    "births",
}

WIKIPEDIA_TITLE_OVERRIDES = {
    base.territory_key("Salcedo", "Hermanas Mirabal"): "Salcedo (República Dominicana)",
}

# These municipalities were created from former municipal districts after the
# statistical and PMD sources used by the portal were produced. Documents from
# the predecessor district are territorial antecedents, not historical PMDs of
# the new municipality.
NEW_MUNICIPALITIES = {
    base.territory_key("Villa Central", "Barahona"),
    base.territory_key("Tireo", "La Vega"),
    base.territory_key("La Caleta", "Santo Domingo"),
    base.territory_key("La Victoria", "Santo Domingo"),
}

COLORS = {
    "ink": "203740",
    "muted": "5E7076",
    "teal": "14866D",
    "blue": "2F6CC6",
    "purple": "7651B8",
    "orange": "C96B21",
    "rose": "C9526A",
    "gold": "A27719",
    "line": "D7E0E1",
    "soft": "F3F7F6",
    "white": "FFFFFF",
    "warn": "9A5A18",
    "warn_soft": "FFF6E7",
}

SERVICE_LABELS = {
    "servicios_sanitarios": {
        "inodoro": "Inodoro",
        "letrina": "Letrina",
        "no_tiene": "No tiene",
        "sin_informacion": "Sin información",
    },
    "agua_uso_domestico": {
        "del_acueducto_dentro_de_la_vivienda": "Acueducto dentro",
        "del_acueducto_en_el_patio_de_la_vivienda": "Acueducto en patio",
        "de_una_llave_publica": "Llave pública",
        "de_una_llave_de_otra_vivienda": "Llave de otra vivienda",
        "de_un_tubo_de_la_calle": "Tubo de la calle",
        "manantial_rio_arroyo": "Manantial / río",
        "pozo_tubular": "Pozo tubular",
        "pozo_cavado": "Pozo cavado",
        "lluvia": "Agua de lluvia",
        "camion_tanque": "Camión tanque",
        "otro": "Otro",
    },
    "agua_para_beber": {
        "del_acueducto_dentro_de_la_vivienda": "Acueducto dentro",
        "del_acueducto_en_el_patio_de_la_vivienda": "Acueducto en patio",
        "de_una_llave_publica": "Llave pública",
        "de_una_llave_de_otra_vivienda": "Llave de otra vivienda",
        "manantial_rio_arroyo": "Manantial / río",
        "pozo_tubular": "Pozo tubular",
        "pozo_cavado": "Pozo cavado",
        "lluvia": "Agua de lluvia",
        "camion_tanque": "Camión tanque",
        "botellones": "Botellones",
        "camioncito_procesada": "Camioncito procesada",
        "otro": "Otro",
    },
    "alumbrado": {
        "energia_electrica_del_tendido_publico": "Red pública",
        "energia_eletrica_del_tendido_publico": "Red pública",
        "lampara_de_gas_propano": "Lámpara de gas",
        "lampara_de_gas_kerosene": "Lámpara de kerosene",
        "energia_electrica_de_planta_propia": "Planta propia",
        "paneles_solares": "Paneles solares",
        "otros": "Otros",
        "sin_informacion": "Sin información",
    },
    "combustible_cocinar": {
        "gas_propano": "Gas propano",
        "carbon": "Carbón",
        "lena": "Leña",
        "electricidad": "Electricidad",
        "otro": "Otro",
        "no_cocina": "No cocina",
        "sin_informacion": "Sin información",
    },
    "eliminacion_basura": {
        "la_recoge_el_ayuntamiento": "Recoge ayuntamiento",
        "la_recoge_una_empresa_privada": "Empresa privada",
        "la_queman": "La queman",
        "la_tiran_en_el_patio_o_sola": "Patio / solar",
        "la_tiran_en_un_vertedero": "Vertedero",
        "la_tiran_en_un_rio_o_canada": "Río / cañada",
        "otros": "Otros",
    },
}


def load_json(path: Path) -> Any:
    with path.open("r", encoding="utf-8") as handle:
        return json.load(handle)


def save_json(path: Path, value: Any) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    with path.open("w", encoding="utf-8", newline="\n") as handle:
        json.dump(value, handle, ensure_ascii=False, indent=2)
        handle.write("\n")


def fmt_number(value: float | int | None, decimals: int = 0) -> str:
    if value is None:
        return "No disponible"
    if isinstance(value, float) and (math.isnan(value) or math.isinf(value)):
        return "No disponible"
    return f"{float(value):,.{decimals}f}"


def fmt_pct(value: float | int | None, decimals: int = 1) -> str:
    if value is None:
        return "No disponible"
    return f"{float(value):.{decimals}f}%"


def fmt_money(value: float | int | None) -> str:
    if value is None:
        return "No disponible"
    amount = float(value)
    if math.isnan(amount) or math.isinf(amount):
        return "No disponible"
    return f"RD$ {amount:,.0f}"


def find_pdftoppm() -> Path:
    candidates = [
        Path.home()
        / ".cache"
        / "codex-runtimes"
        / "codex-primary-runtime"
        / "dependencies"
        / "native"
        / "poppler"
        / "Library"
        / "bin"
        / "pdftoppm.exe",
        Path(shutil.which("pdftoppm") or ""),
        Path(shutil.which("pdftoppm.exe") or ""),
    ]
    for candidate in candidates:
        if str(candidate) and candidate.is_file():
            return candidate
    raise RuntimeError("No se encontró pdftoppm para convertir el PDF del Dashboard.")


def rasterize_dashboard_pdf(
    pdf_path: Path,
    asset_dir: Path,
    *,
    reuse_existing: bool = False,
) -> list[Path]:
    """Render every dashboard PDF page as a readable Word image."""
    asset_dir.mkdir(parents=True, exist_ok=True)
    pattern = "dashboard-pdf-page-*.jpg"
    existing = sorted(
        asset_dir.glob(pattern),
        key=lambda path: int(re.search(r"(\d+)$", path.stem).group(1)),
    )
    if reuse_existing and existing:
        return existing
    for path in existing:
        path.unlink()

    prefix = asset_dir / "dashboard-pdf-page"
    command = [
        str(find_pdftoppm()),
        "-jpeg",
        "-jpegopt",
        "quality=88,optimize=y,progressive=y",
        "-r",
        "144",
        str(pdf_path),
        str(prefix),
    ]
    completed = subprocess.run(
        command,
        check=False,
        capture_output=True,
        text=True,
        encoding="utf-8",
        errors="replace",
    )
    if completed.returncode:
        raise RuntimeError(
            f"No se pudo convertir {pdf_path.name}: "
            f"{(completed.stderr or completed.stdout).strip()}"
        )
    pages = sorted(
        asset_dir.glob(pattern),
        key=lambda path: int(re.search(r"(\d+)$", path.stem).group(1)),
    )
    if not pages:
        raise RuntimeError(f"El PDF {pdf_path.name} no produjo imágenes.")
    return pages


def territorial_indicator_rows(
    municipality: dict[str, Any],
    context: dict[str, Any],
) -> list[list[str]]:
    """Build a scope-safe summary of the new Dashboard Territorial indicators."""
    name = municipality["municipio"]
    province_name = municipality["provincia"]
    municipal = (context.get("municipality") or {}).get("metrics") or {}
    province = (context.get("province") or {}).get("metrics") or {}
    rows: list[list[str]] = []

    def latest(metric: dict[str, Any]) -> dict[str, Any]:
        return metric.get("latest") or {}

    def period(metric: dict[str, Any]) -> str:
        item = latest(metric)
        return str(item.get("period") or item.get("year") or metric.get("year") or "s/f")

    def add(
        label: str,
        value: str,
        scope: str,
        period_value: str,
        source_key: str,
    ) -> None:
        rows.append([label, value, scope, period_value, source_key])

    fires = municipal.get("fires") or {}
    if latest(fires):
        item = latest(fires)
        add(
            "Incendios forestales registrados",
            f"{fmt_number(item.get('records'))} registros; "
            f"{fmt_number(item.get('affected'))} tareas declaradas afectadas",
            f"Municipio de {name}",
            period(fires),
            "Datos abiertos · incendios forestales",
        )

    inaipi = municipal.get("inaipi") or {}
    if inaipi.get("centers") is not None:
        add(
            "Centros INAIPI",
            f"{fmt_number(inaipi.get('centers'))} centros",
            f"Municipio de {name}",
            str(inaipi.get("snapshot") or "s/f"),
            "Datos abiertos · INAIPI",
        )

    sports = municipal.get("sports") or {}
    if sports.get("count") is not None:
        add(
            "Instalaciones deportivas registradas",
            f"{fmt_number(sports.get('count'))} instalaciones; "
            f"{fmt_number(sports.get('per10k'), 1)} por 10 mil habitantes",
            f"Municipio de {name}",
            str(sports.get("year") or "s/f"),
            "Datos abiertos · instalaciones deportivas",
        )

    investment = municipal.get("investment") or {}
    if latest(investment):
        item = latest(investment)
        add(
            "Inversión pública territorial",
            f"Presupuesto {fmt_money(item.get('budget'))}; ejecutado "
            f"{fmt_money(item.get('executed'))}; ejecución {fmt_pct(item.get('executionPct'))}",
            f"Municipio de {name}",
            period(investment),
            "Datos abiertos · estadísticas de proyectos de inversión",
        )

    sismap = municipal.get("sismap") or {}
    if sismap.get("score") is not None:
        add(
            "Puntuación SISMAP Municipal",
            f"{fmt_number(sismap.get('score'), 1)} / 100",
            f"Municipio de {name}",
            str(sismap.get("date") or "s/f"),
            "SISMAP Municipal",
        )

    roads = municipal.get("roads") or {}
    if latest(roads):
        item = latest(roads)
        add(
            "Trabajos viales registrados",
            f"{fmt_number(item.get('records'))} registros; "
            f"{fmt_number(item.get('lengthKm'), 1)} km",
            f"Municipio de {name}",
            period(roads),
            "Datos abiertos · infraestructura vial",
        )

    permits = municipal.get("permits") or {}
    if latest(permits):
        item = latest(permits)
        add(
            "Licencias de construcción",
            f"{fmt_number(item.get('licenses'))} licencias; "
            f"{fmt_number(item.get('areaM2'))} m²; inversión {fmt_money(item.get('investment'))}",
            f"Municipio de {name}",
            period(permits),
            "Datos abiertos · licencias emitidas",
        )

    provincial_scope = f"Provincia completa de {province_name}; no es valor municipal"
    crime = province.get("crime") or {}
    if latest(crime):
        item = latest(crime)
        add(
            "Robos reportados",
            f"{fmt_number(item.get('count'))} casos; "
            f"{fmt_number(item.get('rate'), 1)} por 100 mil habitantes",
            provincial_scope,
            period(crime),
            "Datos abiertos · robos reportados",
        )
    homicide = province.get("homicide") or {}
    if latest(homicide):
        item = latest(homicide)
        add(
            "Homicidios intencionales",
            f"{fmt_number(item.get('count'))} casos; "
            f"{fmt_number(item.get('rate'), 1)} por 100 mil habitantes",
            provincial_scope,
            period(homicide),
            "ONE · muertes accidentales y violentas",
        )
    traffic = province.get("traffic") or {}
    if latest(traffic):
        item = latest(traffic)
        add(
            "Fallecimientos de tránsito",
            f"{fmt_number(item.get('count'))} personas; "
            f"{fmt_number(item.get('rate'), 1)} por 100 mil habitantes",
            provincial_scope,
            period(traffic),
            "Datos abiertos · fallecimientos de tránsito",
        )
    overcrowding = province.get("overcrowding") or {}
    if overcrowding:
        total_pct = float(overcrowding.get("extremePct") or 0) + float(
            overcrowding.get("moderatePct") or 0
        )
        add(
            "Hacinamiento extremo o moderado",
            fmt_pct(total_pct),
            provincial_scope,
            str(overcrowding.get("year") or overcrowding.get("snapshot") or "s/f"),
            "Datos abiertos · hacinamiento",
        )
    disability = province.get("disability") or {}
    if disability:
        vulnerable_pct = float(disability.get("icv1Pct") or 0) + float(
            disability.get("icv2Pct") or 0
        )
        add(
            "Personas con discapacidad en ICV-1 o ICV-2",
            fmt_pct(vulnerable_pct),
            provincial_scope,
            str(disability.get("year") or disability.get("snapshot") or "s/f"),
            "Datos abiertos · discapacidad por provincia",
        )
    births = province.get("births") or {}
    if latest(births):
        item = latest(births)
        add(
            "Registros de nacimiento",
            f"{fmt_number(item.get('count'))} registros; "
            f"{fmt_number(item.get('rate'), 1)} por mil habitantes",
            provincial_scope,
            period(births),
            "Datos abiertos · provincia de registro",
        )
    return rows


def index_mapa_projects(payload: dict[str, Any]) -> dict[str, list[dict[str, Any]]]:
    """Index active/programmed projects by specifically identified municipality."""
    indexed: dict[str, dict[str, dict[str, Any]]] = defaultdict(dict)
    for project in payload.get("projects", []):
        if project.get("state") not in ACTIVE_MAPA_STATES:
            continue
        if project.get("locationScope") != "municipality":
            continue
        project_id = str(project.get("mapProjectId") or project.get("code") or "")
        for location in project.get("locations", []):
            municipality = location.get("municipality")
            province = location.get("province")
            if not municipality or not province:
                continue
            key = base.territory_key(municipality, province)
            indexed[key][project_id] = project
    return {
        key: sorted(
            projects.values(),
            key=lambda item: (
                0 if item.get("state") == "EJECUCIÓN" else 1,
                -float(item.get("budget") or 0),
                str(item.get("name") or ""),
            ),
        )
        for key, projects in indexed.items()
    }


def pct(part: float | int | None, total: float | int | None) -> float | None:
    if part is None or total in (None, 0):
        return None
    return float(part) * 100.0 / float(total)


def sentences(text: str) -> list[str]:
    cleaned = re.sub(r"\s+", " ", text or "").strip()
    cleaned = re.sub(r"\[\s*cita requerida\s*\]", "", cleaned, flags=re.I)
    return [
        item.strip()
        for item in re.split(r"(?<=[.!?])\s+(?=[A-ZÁÉÍÓÚÑ0-9«])", cleaned)
        if 25 <= len(item.strip()) <= 520
    ]


def split_wikipedia_sections(extract: str) -> dict[str, str]:
    parts = re.split(r"\n+==+\s*([^=\n]+?)\s*==+\n+", extract or "")
    result = {"Introducción": parts[0].strip() if parts else ""}
    for index in range(1, len(parts) - 1, 2):
        result[parts[index].strip()] = parts[index + 1].strip()
    return result


def is_municipal_wikipedia(row: dict[str, Any]) -> bool:
    title = str(row.get("title") or "").strip()
    return row.get("status") == "verified" and not base.clean(title).startswith("provinciade")


def sanitize_wikipedia_cache(cache: dict[str, Any], targets: list[dict[str, Any]]) -> None:
    for item in targets:
        key = base.territory_key(item["municipio"], item["provincia"])
        override = WIKIPEDIA_TITLE_OVERRIDES.get(key)
        row = cache.get(key, {})
        if override:
            if row.get("title") != override:
                cache[key] = {
                    "status": "verified",
                    "title": override,
                    "url": "https://es.wikipedia.org/wiki/" + urllib.parse.quote(override.replace(" ", "_")),
                    "extract": "",
                    "retrieved_at": TODAY.isoformat(),
                }
            continue
        if row.get("status") == "verified" and not is_municipal_wikipedia(row):
            cache[key] = {
                "status": "rejected_nonmunicipal",
                "title": row.get("title", ""),
                "url": row.get("url", ""),
                "reason": "La página corresponde a una provincia, no al municipio.",
                "retrieved_at": TODAY.isoformat(),
            }


def fetch_full_wikipedia(cache: dict[str, Any], targets: list[dict[str, Any]]) -> None:
    pending = []
    for item in targets:
        key = base.territory_key(item["municipio"], item["provincia"])
        row = cache.get(key, {})
        if is_municipal_wikipedia(row) and not row.get("full_extract") and row.get("title"):
            pending.append((key, row["title"]))

    # MediaWiki's TextExtracts extension can return an empty extract for all
    # but one page in a multi-title request. Fetch one verified municipal page
    # at a time so a successful lookup is never mistaken for missing content.
    for key, title in pending:
        params = {
            "action": "query",
            "prop": "extracts|info|revisions",
            "explaintext": "1",
            "titles": title,
            "inprop": "url",
            "rvprop": "ids",
            "format": "json",
            "formatversion": "2",
        }
        url = "https://es.wikipedia.org/w/api.php?" + urllib.parse.urlencode(params)
        payload = None
        for attempt in range(3):
            try:
                request = urllib.request.Request(url, headers={"User-Agent": USER_AGENT})
                with urllib.request.urlopen(request, timeout=45) as response:
                    payload = json.load(response)
                break
            except Exception as exc:
                cache.setdefault(key, {})["full_extract_error"] = str(exc)[:180]
                time.sleep(3.0 if "429" in str(exc) else 1.0)
        if payload is None:
            continue

        page = next(iter(payload.get("query", {}).get("pages", [])), None)
        if not page or page.get("missing"):
            continue
        full_extract = page.get("extract", "") or ""
        cache[key]["full_extract"] = full_extract
        cache[key]["sections"] = split_wikipedia_sections(full_extract)
        cache[key]["revid"] = (page.get("revisions") or [{}])[0].get("revid")
        cache[key]["url"] = page.get("fullurl") or cache[key].get("url", "")
        cache[key]["retrieved_at"] = TODAY.isoformat()
        cache[key].pop("full_extract_error", None)
        time.sleep(0.75)


def extract_historical_context(
    rows: list[dict[str, Any]], cache: dict[str, Any], cache_key: str
) -> dict[str, Any]:
    if cache_key in cache:
        cached = cache[cache_key]
        cached["history"] = re.sub(r"(?<=\w)-\s+(?=\w)", "", cached.get("history", ""))
        cached["history"] = re.sub(
            r"(?i)\boriginalmente llamado originalmente\b",
            "Originalmente llamado",
            cached["history"],
        )
        if (
            cached["history"]
            and not re.search(r"(?i)¡?error!?|\.{12,}", cached["history"])
            and len(cached["history"]) >= 120
        ):
            return cached
        cache.pop(cache_key, None)
    result: dict[str, Any] = {"history": "", "history_pages": "", "strategic": "", "strategic_pages": ""}
    candidates = sorted(rows, key=lambda row: row.get("period_end") or 0, reverse=True)
    for row in candidates:
        path = Path(row.get("local_path", ""))
        if not path.exists():
            continue
        if fitz is None:
            continue
        try:
            pdf = fitz.open(path)
        except Exception:
            continue
        page_texts: list[str] = []
        history_hits: list[int] = []
        strategy_hits: list[int] = []
        for index in range(pdf.page_count):
            text = re.sub(r"\s+", " ", pdf[index].get_text("text") or "").strip()
            page_texts.append(text)
            normalized = base.clean(text)
            if any(token in normalized for token in ("antecedenteshistoricos", "resenahistorica", "historiadelmunicipio")):
                history_hits.append(index)
            if "lineasestrategicas" in normalized and len(strategy_hits) < 2:
                strategy_hits.append(index)
        if history_hits and not result["history"]:
            for index in history_hits:
                combined = " ".join(page_texts[index : min(index + 2, len(page_texts))])
                if re.search(r"(?i)¡?error!?|\.{12,}", combined):
                    continue
                combined = re.sub(r"(?<=\w)-\s+(?=\w)", "", combined)
                combined = re.sub(r"(?i)\boriginalmente llamado originalmente\b", "Originalmente llamado", combined)
                start = re.search(
                    r"(?i)(antecedentes\s+hist[oó]ricos|reseña\s+hist[oó]rica|historia\s+del\s+municipio)",
                    combined,
                )
                if start:
                    combined = combined[start.end() :]
                useful = [
                    sentence
                    for sentence in sentences(combined)
                    if not re.match(r"(?i)plan de desarrollo municipal|p[aá]gina \d+", sentence)
                ][:6]
                if len(" ".join(useful)) < 120:
                    continue
                result["history"] = " ".join(useful)[:2400]
                result["history_pages"] = f"{index + 1}-{min(index + 2, len(page_texts))}"
                result["history_period"] = f"{row.get('period_start')}-{row.get('period_end')}"
                result["history_source"] = row.get("source_url") or row.get("source_page") or str(path)
                break
        pdf.close()
        if result["history"]:
            break
    cache[cache_key] = result
    return result


def service_share(living: dict[str, Any] | None, service: str, category: str) -> float | None:
    if not living:
        return None
    item = living.get("servicios", {}).get(service, {})
    categories = item.get("categorias", {})
    value = categories.get(category)
    if value is None and category == "energia_electrica_del_tendido_publico":
        value = categories.get("energia_eletrica_del_tendido_publico")
    return pct(value, item.get("total"))


def national_service_share(national: dict[str, Any], service: str, category: str) -> float | None:
    item = national.get(service, {})
    categories = item.get("categorias", {})
    value = categories.get(category)
    if value is None and category == "energia_electrica_del_tendido_publico":
        value = categories.get("energia_eletrica_del_tendido_publico")
    return pct(value, item.get("total"))


def group_index(rows: list[dict[str, Any]], field: str = "adm2_code") -> dict[str, list[dict[str, Any]]]:
    result: dict[str, list[dict[str, Any]]] = defaultdict(list)
    for row in rows:
        result[base.code_key(row.get(field))].append(row)
    return result


def polygon_rings(geometry: dict[str, Any]) -> Iterable[list[list[float]]]:
    if geometry.get("type") == "Polygon":
        for ring in geometry.get("coordinates", [])[:1]:
            yield ring
    elif geometry.get("type") == "MultiPolygon":
        for polygon in geometry.get("coordinates", []):
            for ring in polygon[:1]:
                yield ring


def dashboard_source_note() -> str:
    return "Fuente: Censo 2022, DEE 2024, Anuario Estadístico Educativo 2024 y registros SNS disponibles en el Dashboard."


def add_card(
    fig,
    rect: tuple[float, float, float, float],
    title: str,
    *,
    accent: str = "#2F6CC6",
    face: str = "#FFFFFF",
    title_size: float = 9.2,
):
    x, y, width, height = rect
    patch = FancyBboxPatch(
        (x, y),
        width,
        height,
        transform=fig.transFigure,
        boxstyle="round,pad=0.006,rounding_size=0.009",
        linewidth=0.7,
        edgecolor="#D7E0E1",
        facecolor=face,
        zorder=0,
    )
    fig.patches.append(patch)
    fig.text(x + 0.015, y + height - 0.029, title, color=accent, fontsize=title_size, weight="bold")
    return (x + 0.018, y + 0.018, width - 0.036, height - 0.065)


def new_dashboard_page(title: str, municipality: str, subtitle: str = ""):
    fig = plt.figure(figsize=(8.5, 10.2), dpi=145, facecolor="#F7F9FA")
    fig.text(0.055, 0.965, title, fontsize=15, weight="bold", color="#203740")
    fig.text(0.055, 0.944, municipality, fontsize=9.5, color="#5E7076")
    if subtitle:
        fig.text(0.945, 0.95, subtitle, fontsize=7.5, color="#7A898D", ha="right")
    return fig


def finish_dashboard_page(fig, path: Path, source: str = "") -> None:
    fig.text(0.055, 0.026, source or dashboard_source_note(), fontsize=6.7, color="#75868B")
    fig.savefig(path, dpi=145, facecolor=fig.get_facecolor(), bbox_inches=None)
    plt.close(fig)


def render_map(
    fig,
    rect: tuple[float, float, float, float],
    geojson: dict[str, Any],
    adm2_code: str,
) -> None:
    ax = fig.add_axes(rect, zorder=2)
    selected = None
    bounds = [999.0, 999.0, -999.0, -999.0]
    for feature in geojson.get("features", []):
        if base.code_key(feature.get("properties", {}).get("adm2_code")) == adm2_code:
            selected = feature
        for ring in polygon_rings(feature.get("geometry", {})):
            xs = [point[0] for point in ring]
            ys = [point[1] for point in ring]
            if not xs:
                continue
            bounds[0] = min(bounds[0], min(xs))
            bounds[1] = min(bounds[1], min(ys))
            bounds[2] = max(bounds[2], max(xs))
            bounds[3] = max(bounds[3], max(ys))
            ax.add_patch(
                Polygon(
                    ring,
                    closed=True,
                    facecolor="#E7ECEC",
                    edgecolor="#FFFFFF",
                    linewidth=0.35,
                )
            )
    if selected:
        for ring in polygon_rings(selected.get("geometry", {})):
            ax.add_patch(
                Polygon(
                    ring,
                    closed=True,
                    facecolor="#E5574F",
                    edgecolor="#A62F2B",
                    linewidth=1.0,
                )
            )
    if bounds[0] < 900:
        ax.set_xlim(bounds[0] - 0.15, bounds[2] + 0.15)
        ax.set_ylim(bounds[1] - 0.08, bounds[3] + 0.08)
    ax.set_aspect("equal")
    ax.axis("off")


def draw_pyramid(ax, records: list[dict[str, Any]], title: str) -> None:
    if not records:
        ax.text(0.5, 0.5, "Datos no disponibles", ha="center", va="center", color="#75868B")
        ax.axis("off")
        return
    rows = [row for row in records if "declarado" not in str(row.get("age_group", "")).lower()]
    labels = [str(row.get("age_group", "")).replace(" - ", "–") for row in rows]
    male = [-float(row.get("male", 0) or 0) for row in rows]
    female = [float(row.get("female", 0) or 0) for row in rows]
    positions = list(range(len(rows)))
    ax.barh(positions, male, color="#2F6CC6", height=0.82)
    ax.barh(positions, female, color="#E94F4F", height=0.82)
    tick_step = 2 if len(labels) > 15 else 1
    ax.set_yticks(positions[::tick_step])
    ax.set_yticklabels(labels[::tick_step], fontsize=5.3)
    max_value = max([abs(value) for value in male + female] or [1])
    ax.set_xlim(-max_value * 1.12, max_value * 1.12)
    ax.axvline(0, color="#839196", linewidth=0.5)
    ax.grid(axis="x", color="#DCE3E5", linewidth=0.45)
    ax.tick_params(axis="x", labelsize=5.2)
    ax.set_title(title, fontsize=7.2, weight="bold", color="#203740", pad=4)
    for side in ax.spines.values():
        side.set_visible(False)


def text_rows(
    fig,
    rect: tuple[float, float, float, float],
    rows: list[tuple[str, str]],
    *,
    color: str = "#203740",
    max_rows: int = 8,
) -> None:
    x, y, width, height = rect
    visible = rows[:max_rows]
    line_height = height / max(1, len(visible))
    for index, (label, value) in enumerate(visible):
        yy = y + height - (index + 0.62) * line_height
        fig.text(x, yy, label, fontsize=6.9, color=color)
        fig.text(x + width, yy, value, fontsize=6.9, color=color, ha="right", weight="bold")


def top_categories(item: dict[str, Any], labels: dict[str, str], limit: int = 6) -> list[tuple[str, str]]:
    total = item.get("total") or 0
    categories = item.get("categorias", {})
    ordered = sorted(categories.items(), key=lambda pair: pair[1] or 0, reverse=True)
    result = []
    for key, value in ordered[:limit]:
        result.append((labels.get(key, key.replace("_", " ").title()), f"{fmt_number(value)} · {fmt_pct(pct(value, total))}"))
    return result


def generate_dashboard_pages(
    municipality: dict[str, Any],
    adm2_code: str,
    data: dict[str, Any],
    national: dict[str, Any],
    geojson: dict[str, Any],
    asset_dir: Path,
    reuse_existing_rest: bool = False,
) -> list[Path]:
    if plt is None:
        raise RuntimeError(
            "Matplotlib no está disponible y no existen láminas del Dashboard para reutilizar."
        )
    asset_dir.mkdir(parents=True, exist_ok=True)
    name = municipality["municipio"]
    pages = [asset_dir / f"diagnostico-{index}.png" for index in range(1, 6)]
    basic = data.get("basic") or {}
    households = data.get("households") or {}
    urban = data.get("urban_rural") or {}
    living = data.get("living") or {}
    pyramid = data.get("pyramid") or []
    pyramid2010 = data.get("pyramid2010") or []
    household_size = data.get("household_size") or []
    education = data.get("education") or {}
    education_level = data.get("education_level") or {}
    education_offer = data.get("education_offer") or {}
    economy = data.get("economy") or {}
    tic = data.get("tic") or {}
    health = data.get("health") or {}

    if not basic:
        gap_pages = [
            ("Diagnóstico municipal · Población y hogares", "Censo 2010 y 2022"),
            ("Diagnóstico municipal · Condición de vida y servicios básicos", "Hogares · Censo 2022"),
            ("Diagnóstico municipal · Educación", "Censo 2022 y Anuario 2024"),
            ("Diagnóstico municipal · Economía y empleo", "DEE 2024"),
            ("Diagnóstico municipal · Salud y comparación", "Registros disponibles"),
        ]
        for index, (title, source_label) in enumerate(gap_pages):
            fig = new_dashboard_page(title, name, source_label)
            content = add_card(
                fig,
                (0.10, 0.25, 0.80, 0.55),
                "Ficha municipal separada aún no disponible",
                accent="#7B8794",
                face="#F7F9F8",
            )
            fig.text(
                content[0] + 0.03,
                content[1] + content[3] * 0.58,
                "El Dashboard no dispone de datos estadísticos\n"
                "correspondientes exactamente al nuevo ámbito municipal.",
                fontsize=15,
                weight="bold",
                color="#203740",
                va="center",
            )
            fig.text(
                content[0] + 0.03,
                content[1] + content[3] * 0.34,
                "No se trasladan cifras del municipio de origen ni promedios\n"
                "provinciales como si fueran datos propios.",
                fontsize=10,
                color="#5E7076",
                va="center",
            )
            finish_dashboard_page(
                fig,
                pages[index],
                "Fuente: verificación de cobertura del Dashboard de Diagnóstico Territorial, 2026.",
            )
        return pages

    # Page 1: demography and households
    fig = new_dashboard_page("Diagnóstico municipal · Población y hogares", name, "Censo 2022")
    cards = [
        ("Población total", fmt_number(basic.get("poblacion_total")), COLORS["blue"]),
        ("Variación 2010–2022", fmt_pct(basic.get("variacion_pct")), COLORS["teal"]),
        (
            "Hombres / mujeres",
            f"{fmt_pct(pct(basic.get('poblacion_hombres'), basic.get('poblacion_total')))} / "
            f"{fmt_pct(pct(basic.get('poblacion_mujeres'), basic.get('poblacion_total')))}",
            COLORS["purple"],
        ),
        (
            "Urbana / rural",
            f"{fmt_pct(pct(urban.get('urbana'), urban.get('poblacion_total')))} / "
            f"{fmt_pct(pct(urban.get('rural'), urban.get('poblacion_total')))}",
            COLORS["orange"],
        ),
    ]
    for index, (label, value, accent) in enumerate(cards):
        x = 0.055 + index * 0.225
        content = add_card(fig, (x, 0.835, 0.205, 0.082), label, accent=f"#{accent}", face="#FFFFFF")
        fig.text(content[0], content[1] + 0.018, value, fontsize=12, weight="bold", color="#203740")
    map_rect = add_card(fig, (0.055, 0.50, 0.43, 0.31), "Ubicación geográfica", accent="#2F6CC6")
    render_map(fig, map_rect, geojson, adm2_code)
    pyramid_rect = add_card(fig, (0.505, 0.50, 0.44, 0.31), "Pirámide de población 2022", accent="#2F6CC6")
    draw_pyramid(fig.add_axes(pyramid_rect, zorder=2), pyramid, "")
    household_rect = add_card(fig, (0.055, 0.12, 0.43, 0.35), "Hogares y viviendas", accent="#2F6CC6")
    text_rows(
        fig,
        household_rect,
        [
            ("Hogares totales", fmt_number(households.get("hogares_total"))),
            ("Población en hogares", fmt_number(households.get("poblacion_en_hogares"))),
            ("Personas por hogar", fmt_number(households.get("personas_por_hogar"), 2)),
            ("Viviendas totales", fmt_number(basic.get("viviendas_total"))),
            ("Viviendas ocupadas", fmt_number(basic.get("viviendas_ocupadas"))),
            ("Viviendas desocupadas", fmt_number(basic.get("viviendas_desocupadas"))),
        ],
    )
    pyramid_2010_rect = add_card(
        fig,
        (0.505, 0.12, 0.44, 0.35),
        "Pirámide de población 2010",
        accent="#7B8794",
    )
    draw_pyramid(fig.add_axes(pyramid_2010_rect, zorder=2), pyramid2010, "")
    finish_dashboard_page(
        fig,
        pages[0],
        "Fuente: IX y X Censos Nacionales de Población y Vivienda 2010 y 2022, ONE.",
    )
    if reuse_existing_rest and all(path.exists() for path in pages[1:]):
        return pages

    # Page 2: living conditions
    fig = new_dashboard_page("Diagnóstico municipal · Condición de vida y servicios básicos", name, "Hogares · Censo 2022")
    service_specs = [
        ("Servicios sanitarios", "servicios_sanitarios", COLORS["teal"]),
        ("Agua para uso doméstico", "agua_uso_domestico", COLORS["teal"]),
        ("Agua para beber", "agua_para_beber", COLORS["teal"]),
        ("Alumbrado principal", "alumbrado", COLORS["orange"]),
        ("Combustible para cocinar", "combustible_cocinar", COLORS["orange"]),
        ("Eliminación de basura", "eliminacion_basura", COLORS["orange"]),
    ]
    for index, (title, key, accent) in enumerate(service_specs):
        col = index % 3
        row = index // 3
        rect = (0.055 + col * 0.300, 0.535 - row * 0.405, 0.28, 0.375)
        content = add_card(fig, rect, title, accent=f"#{accent}", face="#FFFFFF")
        item = (living.get("servicios") or {}).get(key, {})
        text_rows(fig, content, top_categories(item, SERVICE_LABELS[key], 7), max_rows=7)
    finish_dashboard_page(fig, pages[1], "Fuente: X Censo Nacional de Población y Vivienda 2022, ONE.")

    # Page 3: education
    fig = new_dashboard_page("Diagnóstico municipal · Educación", name, "Censo 2022 y Anuario 2024")
    offer_rect = add_card(fig, (0.055, 0.62, 0.40, 0.29), "Oferta educativa asociada", accent="#C96B21", face="#FFF9F1")
    offer_rows = [("Centros totales", fmt_number(education_offer.get("centros_total")))]
    for level, label in (("inicial_primario", "Inicial / Primario"), ("secundario", "Secundario"), ("adultos", "Adultos")):
        row = (education_offer.get("niveles") or {}).get(level, {})
        offer_rows.append((label, f"{fmt_number(row.get('centros'))} centros · {fmt_number(row.get('matricula'))} estudiantes"))
    text_rows(fig, offer_rect, offer_rows, max_rows=5)
    level_rect = add_card(fig, (0.475, 0.62, 0.47, 0.29), "Nivel de instrucción", accent="#C96B21", face="#FFF9F1")
    ax = fig.add_axes((level_rect[0], level_rect[1], level_rect[2] * 0.45, level_rect[3]), zorder=2)
    levels = education_level.get("nivel") or education_level.get("niveles") or {}
    level_order = ["ninguno", "preprimaria", "primaria", "primaria_basica", "secundaria", "secundaria_media", "superior", "universitaria_superior", "postgrado"]
    values = []
    labels = []
    seen = set()
    label_map = {
        "ninguno": "Ninguno",
        "preprimaria": "Preprimaria",
        "primaria": "Primaria",
        "primaria_basica": "Primaria",
        "secundaria": "Secundaria",
        "secundaria_media": "Secundaria",
        "superior": "Superior",
        "universitaria_superior": "Superior",
        "postgrado": "Postgrado",
    }
    for key in level_order:
        label = label_map[key]
        if key in levels and label not in seen:
            seen.add(label)
            labels.append(label)
            values.append(float((levels[key] or {}).get("total", 0) or 0))
    if sum(values):
        ax.pie(values, colors=["#EF4444", "#F59E0B", "#22C55E", "#38BDF8", "#8B5CF6"][: len(values)], startangle=90)
        ax.axis("equal")
    else:
        ax.text(0.5, 0.5, "Sin datos", ha="center", va="center")
        ax.axis("off")
    total_levels = sum(values)
    text_rows(
        fig,
        (level_rect[0] + level_rect[2] * 0.50, level_rect[1], level_rect[2] * 0.50, level_rect[3]),
        [(label, fmt_pct(pct(value, total_levels))) for label, value in zip(labels, values)],
        max_rows=6,
    )
    infra_rect = add_card(fig, (0.055, 0.33, 0.89, 0.25), "Infraestructura educativa del distrito asociado", accent="#C96B21", face="#FFF9F1")
    infra = ((education.get("anuario") or {}).get("infraestructura") or {})
    infra_rows = [
        ("Aulas por plantel", fmt_number(infra.get("aulas_por_plantel"), 1)),
        ("Secciones por centro", fmt_number(infra.get("secciones_por_centro"), 1)),
        ("Alumnos por aula", fmt_number(infra.get("alumnos_por_aula"), 1)),
        ("Alumnos por sección", fmt_number(infra.get("alumnos_por_seccion"), 1)),
        ("Alumnos por docente", fmt_number(infra.get("alumnos_por_docente"), 1)),
        ("Docentes por centro", fmt_number(infra.get("docentes_por_centro"), 1)),
    ]
    for index, (label, value) in enumerate(infra_rows):
        col = index % 3
        row = index // 3
        x = infra_rect[0] + col * infra_rect[2] / 3
        y = infra_rect[1] + infra_rect[3] - (row + 1) * infra_rect[3] / 2
        fig.text(x, y + 0.042, label, fontsize=6.7, color="#6C5A4C")
        fig.text(x, y + 0.015, value, fontsize=11, weight="bold", color="#203740")
    eff_rect = add_card(fig, (0.055, 0.08, 0.89, 0.21), "Eficiencia del sistema educativo", accent="#C96B21", face="#FFF9F1")
    efficiency = ((education.get("anuario") or {}).get("eficiencia") or {})
    for index, level in enumerate(("inicial", "primario", "secundario")):
        values_eff = efficiency.get(level) or {}
        x = eff_rect[0] + index * eff_rect[2] / 3
        fig.text(x, eff_rect[1] + eff_rect[3] - 0.030, level.capitalize(), fontsize=7.3, weight="bold", color="#C96B21")
        text_rows(
            fig,
            (x, eff_rect[1], eff_rect[2] / 3 - 0.025, eff_rect[3] - 0.045),
            [
                ("Promoción", fmt_pct(values_eff.get("promocion"))),
                ("Abandono", fmt_pct(values_eff.get("abandono"))),
                ("Reprobación", fmt_pct(values_eff.get("reprobacion"))),
            ],
            max_rows=3,
        )
    finish_dashboard_page(fig, pages[2], "Fuente: MINERD, Anuario Estadístico Educativo 2024 y Censo 2022.")

    # Page 4: economy and employment
    fig = new_dashboard_page("Diagnóstico municipal · Economía y empleo", name, "DEE 2024")
    dee = economy.get("dee_2024") or {}
    econ_cards = [
        ("Establecimientos", fmt_number(dee.get("total_establishments"))),
        ("Empleo estimado", fmt_number(dee.get("total_employees"), 1)),
        ("Empleo medio", fmt_number(dee.get("avg_employees_per_establishment"), 1)),
        (
            "Estab. / 1,000 hab.",
            fmt_number(
                (float(dee.get("total_establishments")) * 1000 / float(basic.get("poblacion_total")))
                if dee.get("total_establishments") is not None and basic.get("poblacion_total")
                else None,
                1,
            ),
        ),
    ]
    for index, (label, value) in enumerate(econ_cards):
        x = 0.055 + index * 0.225
        content = add_card(fig, (x, 0.82, 0.205, 0.09), label, accent="#C9526A", face="#FFF5F7")
        fig.text(content[0], content[1] + 0.018, value, fontsize=11.5, weight="bold", color="#7D3241")
    size_rect = add_card(fig, (0.055, 0.48, 0.41, 0.30), "Empleo por tamaño de establecimiento", accent="#C9526A", face="#FFF5F7")
    bands = dee.get("employment_size_bands") or []
    ax = fig.add_axes((size_rect[0], size_rect[1], size_rect[2] * 0.48, size_rect[3]), zorder=2)
    band_values = [float(row.get("employees", 0) or 0) for row in bands]
    if sum(band_values):
        ax.pie(band_values, colors=["#0EA5E9", "#22C55E", "#F97316", "#8B5CF6"], startangle=90)
        ax.axis("equal")
    else:
        ax.text(0.5, 0.5, "Sin datos", ha="center", va="center")
        ax.axis("off")
    text_rows(
        fig,
        (size_rect[0] + size_rect[2] * 0.51, size_rect[1], size_rect[2] * 0.49, size_rect[3]),
        [(row.get("label", ""), fmt_pct(float(row.get("employees_share", 0) or 0) * 100)) for row in bands],
        max_rows=5,
    )
    sector_rect = add_card(fig, (0.49, 0.48, 0.455, 0.30), "Principales secciones CIIU por empleo", accent="#C9526A", face="#FFF5F7")
    sectors = dee.get("sectors") or []
    sector_rows = [
        (
            textwrap.shorten(str(row.get("label", "")), width=43, placeholder="…"),
            f"{fmt_number(row.get('employees'), 1)} · LQ {fmt_number(row.get('lq'), 2)}",
        )
        for row in sectors[:6]
    ]
    text_rows(fig, sector_rect, sector_rows, max_rows=6)
    bar_rect = add_card(fig, (0.055, 0.10, 0.89, 0.34), "Número de establecimientos por tamaño", accent="#C9526A", face="#FFF5F7")
    ax = fig.add_axes(bar_rect, zorder=2)
    if bands:
        labels = [row.get("label", "") for row in bands]
        establishments = [float(row.get("establishments", 0) or 0) for row in bands]
        ax.bar(range(len(labels)), establishments, color=["#0EA5E9", "#22C55E", "#F97316", "#8B5CF6"])
        ax.set_xticks(range(len(labels)))
        ax.set_xticklabels(labels, fontsize=6.5)
        ax.tick_params(axis="y", labelsize=6)
        ax.grid(axis="y", color="#E5DDE0", linewidth=0.5)
        for side in ax.spines.values():
            side.set_visible(False)
    else:
        ax.text(0.5, 0.5, "Datos no disponibles", ha="center", va="center", color="#75868B")
        ax.axis("off")
    finish_dashboard_page(fig, pages[3], "Fuente: Directorio de Empresas y Establecimientos (DEE) 2024.")

    # Page 5: health and comparison
    fig = new_dashboard_page("Diagnóstico municipal · Salud y comparación", name, "Registros disponibles")
    centers = health.get("centros") or []
    by_type = Counter((row.get("tipo_centro") or "Sin clasificación").title() for row in centers)
    health_rect = add_card(fig, (0.055, 0.57, 0.89, 0.34), f"Establecimientos de salud registrados · Total {len(centers)}", accent="#A27719", face="#FFFBEF")
    ax = fig.add_axes(health_rect, zorder=2)
    if by_type:
        items = by_type.most_common(9)
        labels = [textwrap.shorten(label, width=35, placeholder="…") for label, _ in items]
        values = [value for _, value in items]
        ax.barh(range(len(labels)), values, color="#D29A31")
        ax.set_yticks(range(len(labels)))
        ax.set_yticklabels(labels, fontsize=6.2)
        ax.invert_yaxis()
        ax.tick_params(axis="x", labelsize=6)
        ax.grid(axis="x", color="#E8E1CF", linewidth=0.5)
        for side in ax.spines.values():
            side.set_visible(False)
    else:
        ax.text(0.5, 0.5, "No se dispone de registros municipales en el Dashboard", ha="center", va="center", color="#75868B")
        ax.axis("off")
    compare_rect = add_card(fig, (0.055, 0.10, 0.89, 0.42), "Resumen comparativo: municipio y República Dominicana", accent="#2F6CC6")
    internet_local = ((tic.get("internet") or {}).get("rate_used"))
    computer_local = ((tic.get("computer") or {}).get("rate_used"))
    national_tic = national.get("tic") or {}
    compare = [
        ("Población 2022", fmt_number(basic.get("poblacion_total")), fmt_number((national.get("basic") or {}).get("poblacion_total"))),
        ("Personas por hogar", fmt_number(households.get("personas_por_hogar"), 2), fmt_number((national.get("hogares") or {}).get("personas_por_hogar"), 2)),
        (
            "Agua de acueducto dentro",
            fmt_pct(service_share(living, "agua_uso_domestico", "del_acueducto_dentro_de_la_vivienda")),
            fmt_pct(national_service_share(national.get("living") or {}, "agua_uso_domestico", "del_acueducto_dentro_de_la_vivienda")),
        ),
        (
            "Hogares con inodoro",
            fmt_pct(service_share(living, "servicios_sanitarios", "inodoro")),
            fmt_pct(national_service_share(national.get("living") or {}, "servicios_sanitarios", "inodoro")),
        ),
        (
            "Recogida municipal de basura",
            fmt_pct(service_share(living, "eliminacion_basura", "la_recoge_el_ayuntamiento")),
            fmt_pct(national_service_share(national.get("living") or {}, "eliminacion_basura", "la_recoge_el_ayuntamiento")),
        ),
        (
            "Uso de internet",
            fmt_pct(float(internet_local) * 100 if internet_local is not None else None),
            fmt_pct(float((national_tic.get("internet") or {}).get("rate_used")) * 100 if (national_tic.get("internet") or {}).get("rate_used") is not None else None),
        ),
        (
            "Uso de computadora",
            fmt_pct(float(computer_local) * 100 if computer_local is not None else None),
            fmt_pct(float((national_tic.get("computer") or {}).get("rate_used")) * 100 if (national_tic.get("computer") or {}).get("rate_used") is not None else None),
        ),
        (
            "Centros de salud / 10,000 hab.",
            fmt_number(len(centers) * 10000 / basic.get("poblacion_total") if basic.get("poblacion_total") else None, 2),
            fmt_number(
                (national.get("health") or {}).get("total_centros", 0) * 10000
                / (national.get("basic") or {}).get("poblacion_total", 1),
                2,
            ),
        ),
    ]
    x0, y0, width, height = compare_rect
    fig.text(x0, y0 + height - 0.028, "Indicador", fontsize=7.2, weight="bold", color="#203740")
    fig.text(x0 + width * 0.72, y0 + height - 0.028, name, fontsize=7.2, weight="bold", color="#2F6CC6", ha="right")
    fig.text(x0 + width, y0 + height - 0.028, "País", fontsize=7.2, weight="bold", color="#5E7076", ha="right")
    line_height = (height - 0.050) / len(compare)
    for index, (label, local_value, national_value) in enumerate(compare):
        yy = y0 + height - 0.055 - index * line_height
        fig.text(x0, yy, label, fontsize=6.7, color="#203740")
        fig.text(x0 + width * 0.72, yy, local_value, fontsize=6.7, color="#2F6CC6", ha="right", weight="bold")
        fig.text(x0 + width, yy, national_value, fontsize=6.7, color="#5E7076", ha="right")
    finish_dashboard_page(fig, pages[4], dashboard_source_note())
    return pages


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.find(qn("w:tcMar"))
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa: list[int]) -> None:
    total = sum(widths_dxa)
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            cell.width = Inches(widths_dxa[index] / 1440)
            tc_w = cell._tc.get_or_add_tcPr().find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                cell._tc.get_or_add_tcPr().append(tc_w)
            tc_w.set(qn("w:w"), str(widths_dxa[index]))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)


def set_table_borders(table, color: str = "D7E0E1", size: int = 5) -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = borders.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), str(size))
        node.set(qn("w:color"), color)


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def add_run(paragraph, text: str, *, bold=False, italic=False, color: str | None = None, size: float | None = None):
    run = paragraph.add_run(text)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    if size:
        run.font.size = Pt(size)
    run.font.name = "Calibri"
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Calibri")
    return run


def add_source_note(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph(style="Source Note")
    add_run(paragraph, f"Fuente: {text}", color=COLORS["muted"], size=8.5)


def add_callout(doc: Document, title: str, body: str, *, warning: bool = False) -> None:
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [9360])
    set_table_borders(table, "E3D2B7" if warning else "C9DCD7", 5)
    cell = table.cell(0, 0)
    set_cell_shading(cell, COLORS["warn_soft"] if warning else COLORS["soft"])
    set_cell_margins(cell, top=150, start=180, bottom=150, end=180)
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(3)
    add_run(
        paragraph,
        title,
        bold=True,
        color=COLORS["warn"] if warning else COLORS["teal"],
        size=10.5,
    )
    paragraph = cell.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(0)
    add_run(paragraph, body, color=COLORS["ink"], size=9.5)
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(0)


def add_data_table(
    doc: Document,
    headers: list[str],
    rows: Iterable[Iterable[Any]],
    widths_dxa: list[int],
) -> None:
    values = [list(row) for row in rows]
    if sum(widths_dxa) != 9360:
        raise ValueError(f"Table widths must total 9360 DXA: {widths_dxa}")
    table = doc.add_table(rows=1, cols=len(headers))
    set_table_geometry(table, widths_dxa)
    set_table_borders(table)
    header = table.rows[0]
    set_repeat_table_header(header)
    for index, value in enumerate(headers):
        cell = header.cells[index]
        set_cell_shading(cell, COLORS["ink"])
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        paragraph = cell.paragraphs[0]
        paragraph.paragraph_format.space_after = Pt(0)
        add_run(paragraph, str(value), bold=True, color=COLORS["white"], size=8.7)
    for row_index, row_values in enumerate(values):
        row = table.add_row()
        for index, value in enumerate(row_values):
            cell = row.cells[index]
            if row_index % 2:
                set_cell_shading(cell, "F7F9F9")
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            paragraph = cell.paragraphs[0]
            paragraph.paragraph_format.space_after = Pt(0)
            add_run(paragraph, str(value), color=COLORS["ink"], size=8.7)
    set_table_geometry(table, widths_dxa)
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(0)


def add_page_break(doc: Document) -> None:
    paragraph = doc.add_paragraph()
    paragraph.add_run().add_break(WD_BREAK.PAGE)


def configure_document(doc: Document, municipality: str, province: str, code: str) -> None:
    section = doc.sections[0]
    section.orientation = WD_ORIENT.PORTRAIT
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(COLORS["ink"])
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.333

    style_tokens = {
        "Heading 1": (16, "2E74B5", 18, 10),
        "Heading 2": (13, "2E74B5", 12, 6),
        "Heading 3": (12, "1F4D78", 8, 4),
    }
    for name, (size, color, before, after) in style_tokens.items():
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for name in ("List Bullet", "List Number"):
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
        style.font.size = Pt(11)
        style.paragraph_format.left_indent = Inches(0.375)
        style.paragraph_format.first_line_indent = Inches(-0.194)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.208

    if "Source Note" not in styles:
        style = styles.add_style("Source Note", 1)
    else:
        style = styles["Source Note"]
    style.font.name = "Calibri"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    style.font.size = Pt(8.5)
    style.font.color.rgb = RGBColor.from_string(COLORS["muted"])
    style.paragraph_format.space_before = Pt(4)
    style.paragraph_format.space_after = Pt(4)
    style.paragraph_format.line_spacing = 1.0

    header = section.header
    header_table = header.add_table(rows=1, cols=2, width=Inches(6.5))
    set_table_geometry(header_table, [6000, 3360])
    for cell in header_table.row_cells(0):
        set_cell_margins(cell, top=0, start=0, bottom=0, end=0)
    left = header_table.cell(0, 0).paragraphs[0]
    add_run(left, "PLAN MUNICIPAL DE DESARROLLO", bold=True, color=COLORS["teal"], size=7.5)
    right = header_table.cell(0, 1).paragraphs[0]
    right.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    add_run(right, f"{municipality} · {code or 'código por confirmar'}", color=COLORS["muted"], size=7.5)

    footer = section.footer
    paragraph = footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(paragraph, f"{municipality}, {province} · PMD {PERIOD} · ", color=COLORS["muted"], size=7.5)
    field = OxmlElement("w:fldSimple")
    field.set(qn("w:instr"), "PAGE")
    paragraph._p.append(field)


def age_structure(pyramid: list[dict[str, Any]]) -> dict[str, float | None]:
    if not pyramid:
        return {"youth": None, "working": None, "older": None, "dependency": None}
    youth = working = older = total = 0
    for row in pyramid:
        label = str(row.get("age_group", "")).strip().lower()
        value = float(row.get("male", 0) or 0) + float(row.get("female", 0) or 0)
        if "declarado" in label:
            continue
        total += value
        first = 0 if "menos" in label else int(re.search(r"\d+", label).group()) if re.search(r"\d+", label) else 0
        if first <= 14:
            youth += value
        elif first >= 65:
            older += value
        else:
            working += value
    dependency = (youth + older) * 100 / working if working else None
    return {
        "youth": pct(youth, total),
        "working": pct(working, total),
        "older": pct(older, total),
        "dependency": dependency,
    }


def paragraph(doc: Document, text: str) -> None:
    if text:
        doc.add_paragraph(text)


def write_general_information(
    doc: Document,
    municipality: dict[str, Any],
    code: str,
    data: dict[str, Any],
    wikipedia: dict[str, Any],
    historical: list[dict[str, Any]],
    territorial_antecedents: list[dict[str, Any]],
    historical_context: dict[str, Any],
    area_km2: float | None,
) -> None:
    name = municipality["municipio"]
    province = municipality["provincia"]
    region = municipality["region"]
    basic = data.get("basic") or {}
    households = data.get("households") or {}
    urban = data.get("urban_rural") or {}
    population = basic.get("poblacion_total")
    density = float(population) / float(area_km2) if population and area_km2 else None

    doc.add_heading("1. Información general", level=1)
    doc.add_heading("1.1 Equipo de elaboración y participación del CDM", level=2)
    paragraph(
        doc,
        "La elaboración y gestión del PMD se organiza por responsabilidades institucionales. "
        "Esta forma de presentación mantiene claras las funciones durante todo el período del plan, "
        "aun cuando cambien las personas que ocupan los cargos. La representación comunitaria y "
        "sectorial se canaliza mediante el Consejo de Desarrollo Municipal (CDM), conforme a la Ley núm. 176-07.",
    )
    add_data_table(
        doc,
        ["Instancia", "Responsabilidad en el PMD"],
        [
            ["Alcaldía", "Conducir políticamente el proceso, articular instituciones y presentar la programación anual vinculada al plan."],
            ["Concejo de Regidores", "Conocer, deliberar, aprobar y fiscalizar los instrumentos, presupuestos y ordenanzas que correspondan."],
            ["Oficina Municipal de Planificación y Programación (OMPP)", "Coordinar el diagnóstico, la formulación, la programación, los indicadores y los informes de avance."],
            ["Consejo de Desarrollo Municipal (CDM)", "Participar en la elaboración, discusión y seguimiento del plan; canalizar prioridades territoriales y sectoriales."],
            ["Áreas técnicas municipales", "Aportar información y ejecutar acciones dentro de sus competencias."],
            ["Representación social y económica", "Aportar conocimiento local, contrastar prioridades y acompañar el seguimiento mediante el CDM."],
        ],
        [2900, 6460],
    )
    add_source_note(
        doc,
        "Ley núm. 176-07 del Distrito Nacional y los Municipios, arts. 123–125; "
        "MEPyD, Guía para la formulación de planes de desarrollo municipales, pp. 21–22.",
    )

    doc.add_heading("1.2 Introducción y ¿qué es un PMD?", level=2)
    paragraph(
        doc,
        "El Plan Municipal de Desarrollo (PMD) es el instrumento que orienta las decisiones del "
        "ayuntamiento durante un período de cuatro años. Organiza una visión compartida del territorio, "
        "identifica sus condiciones y necesidades, define objetivos y convierte las prioridades en "
        "acciones que puedan relacionarse con el presupuesto, la gestión de servicios y la coordinación con el Gobierno central.",
    )
    paragraph(
        doc,
        "El PMD no sustituye el presupuesto ni los planes operativos anuales: los articula. Cuando una "
        "necesidad corresponde a agua potable, salud, educación, carreteras u otra institución sectorial, "
        "el plan puede incorporarla como agenda de gestión y coordinación, con resultados e indicadores verificables.",
    )
    add_source_note(
        doc,
        "MEPyD, Guía para la formulación de planes de desarrollo municipales, pp. 19–21 y 46–49; "
        "Ley núm. 176-07, arts. 122–125.",
    )

    doc.add_heading("1.3 Identidad territorial, historia y configuración municipal", level=2)
    paragraph(
        doc,
        f"{name} forma parte de la provincia {province} y de la región de planificación {region}. "
        f"El código geográfico empleado es {code or 'pendiente de confirmación'}. "
        "La identificación territorial enlaza el clasificador geográfico, la cartografía, el Dashboard "
        "de Diagnóstico Territorial y las fuentes documentales. [GEO-01]",
    )
    add_data_table(
        doc,
        ["Dato territorial", "Valor", "Fuente / año"],
        [
            ["Municipio", name, "Portal municipal · 2026"],
            ["Provincia", province, "Portal municipal · 2026"],
            ["Región", region, "Clasificador geográfico"],
            ["Código geográfico", code or "Por confirmar", "Clasificador geográfico"],
            ["Superficie cartográfica", f"{fmt_number(area_km2, 1)} km²" if area_km2 else "No disponible", "GeoJSON municipal"],
            ["Densidad estimada", f"{fmt_number(density, 1)} hab./km²" if density else "No disponible", "Censo 2022 + GeoJSON"],
        ],
        [2600, 3000, 3760],
    )

    history_written = False
    if historical_context.get("history"):
        history_sentences = sentences(historical_context["history"])[:5]
        if history_sentences:
            context_label = (
                "El plan del antiguo distrito municipal, utilizado únicamente como antecedente territorial, "
                if not historical and territorial_antecedents
                else f"El PMD anterior del período {historical_context.get('history_period', 'no identificado')} "
            )
            paragraph(
                doc,
                context_label
                + "recoge los siguientes antecedentes sobre la formación del territorio. "
                + " ".join(history_sentences)
                + f" [PMD-ANT, págs. {historical_context.get('history_pages')}]",
            )
            history_written = True
    wiki_sections = wikipedia.get("sections") or split_wikipedia_sections(wikipedia.get("full_extract", ""))
    wiki_history = ""
    if is_municipal_wikipedia(wikipedia) and not history_written:
        preferred_keys = ("historia", "origen", "formacion", "etimologia", "toponimia")
        wiki_history = next(
            (
                value
                for key, value in wiki_sections.items()
                if any(token in base.clean(key) for token in preferred_keys)
            ),
            "",
        )
        if not wiki_history:
            wiki_history = wiki_sections.get("Introducción", "")
    if wiki_history and not history_written:
        useful = sentences(wiki_history)[:6]
        if useful:
            paragraph(
                doc,
                "Como referencia secundaria, la ficha municipal de Wikipedia resume la identidad y evolución "
                "histórica local de la siguiente manera: "
                + " ".join(useful)
                + " [WIKI-01].",
            )
            history_written = True
    if not history_written:
        paragraph(
            doc,
            "No se incorporó una reseña histórica detallada porque las fuentes disponibles no permitieron "
            "identificar un texto verificable. La OMPP debe completar este apartado con la norma de creación, "
            "el archivo municipal o una publicación histórica institucional.",
        )
    add_source_note(
        doc,
        "Wikipedia se utiliza como apoyo secundario para historia y geografía. Los documentos anteriores "
        "se citan con período y páginas; para municipios nuevos, un plan del antiguo distrito no se presenta "
        "como PMD histórico del municipio actual.",
    )

    doc.add_heading("1.4 Perfil municipal en cifras", level=2)
    if population:
        male = basic.get("poblacion_hombres")
        female = basic.get("poblacion_mujeres")
        variation = basic.get("variacion_pct")
        direction = "creció" if (variation or 0) >= 0 else "disminuyó"
        paragraph(
            doc,
            f"En 2022, {name} registró {fmt_number(population)} habitantes: {fmt_number(male)} hombres "
            f"({fmt_pct(pct(male, population))}) y {fmt_number(female)} mujeres "
            f"({fmt_pct(pct(female, population))}). Frente a los {fmt_number(basic.get('poblacion_2010'))} "
            f"habitantes censados en 2010, la población {direction} {fmt_pct(abs(float(variation)) if variation is not None else None)}. "
            "El cambio intercensal constituye la línea base para dimensionar servicios, infraestructura y demanda de suelo; "
            "no permite atribuir por sí solo causas de migración, fecundidad o movilidad residencial. [DASH-BASE]",
        )
        urban_total = urban.get("urbana")
        rural_total = urban.get("rural")
        paragraph(
            doc,
            f"La distribución territorial muestra {fmt_number(urban_total)} habitantes en zona urbana "
            f"({fmt_pct(pct(urban_total, urban.get('poblacion_total')))}) y {fmt_number(rural_total)} en zona rural "
            f"({fmt_pct(pct(rural_total, urban.get('poblacion_total')))}). Esta relación debe considerarse al organizar "
            "rutas de servicios, mantenimiento vial, equipamientos y mecanismos de atención a comunidades dispersas. [DASH-URB]",
        )
        paragraph(
            doc,
            f"El municipio cuenta con {fmt_number(households.get('hogares_total'))} hogares y "
            f"{fmt_number(households.get('poblacion_en_hogares'))} personas residentes en hogares. "
            f"El promedio es de {fmt_number(households.get('personas_por_hogar'), 2)} personas por hogar. "
            f"El Censo registra {fmt_number(basic.get('viviendas_total'))} viviendas, de las cuales "
            f"{fmt_number(basic.get('viviendas_ocupadas'))} estaban ocupadas y "
            f"{fmt_number(basic.get('viviendas_desocupadas'))} desocupadas. [DASH-HOG] [DASH-BASE]",
        )
    else:
        paragraph(
            doc,
            "El Dashboard aún no dispone de una ficha estadística municipal separada para este territorio. "
            "Por ello no se trasladaron cifras del municipio de origen ni de la provincia. La OMPP debe incorporar "
            "una línea base oficial que corresponda exactamente al nuevo ámbito municipal.",
        )

    doc.add_heading("1.5 Antecedentes de planificación", level=2)
    if historical:
        add_data_table(
            doc,
            ["PMD localizado", "Período", "Condición", "Uso permitido"],
            [
                [
                    f"Plan Municipal de Desarrollo de {name}",
                    f"{row.get('period_start')}-{row.get('period_end')}",
                    "Oficial" if row.get("document_kind") == "official" else "Borrador",
                    "Antecedente histórico; verificar continuidad",
                ]
                for row in sorted(historical, key=lambda item: item.get("period_end") or 0, reverse=True)
            ],
            [3000, 1400, 1500, 3460],
        )
        paragraph(
            doc,
            "Los PMD anteriores permiten reconstruir la trayectoria de problemas, activos, objetivos y proyectos. "
            "En este borrador se utilizan para contextualizar la historia local y señalar temas de continuidad. "
            "No se presume que una obra esté ejecutada, que un problema permanezca vigente ni que una prioridad anterior "
            "haya sido ratificada para 2025–2028 sin verificación de la OMPP. [PMD-ANT]",
        )
    else:
        paragraph(
            doc,
            "No se localizó un PMD anterior del municipio actual en el inventario descargado. "
            "Esta ausencia en el repositorio no demuestra que el documento no exista.",
        )
    if territorial_antecedents:
        add_data_table(
            doc,
            ["Antecedente territorial", "Período", "Alcance correcto", "Uso en este documento"],
            [
                [
                    f"Plan del antiguo Distrito Municipal de {name}",
                    f"{row.get('period_start')}-{row.get('period_end')}",
                    "Documento anterior a la creación del municipio",
                    "Historia y configuración territorial; no se cuenta como PMD del municipio actual",
                ]
                for row in sorted(
                    territorial_antecedents,
                    key=lambda item: item.get("period_end") or 0,
                    reverse=True,
                )
            ],
            [2800, 1200, 2500, 2860],
        )

    add_page_break(doc)
    doc.add_heading("1.6 Marco jurídico", level=2)
    paragraph(
        doc,
        "El PMD se formula dentro del sistema jurídico dominicano de planificación y gestión territorial. "
        "Su aplicación exige coherencia entre el plan, el presupuesto municipal, los instrumentos de "
        "ordenamiento, las competencias sectoriales y los mecanismos de participación y control social.",
    )
    add_data_table(
        doc,
        ["Norma", "Aplicación al Plan Municipal de Desarrollo"],
        [
            ["Constitución de la República Dominicana", "Reconoce la autonomía municipal y la planificación del desarrollo económico y social."],
            ["Ley núm. 176-07 del Distrito Nacional y los Municipios", "Define finalidad, elaboración, participación del CDM, coordinación de la OMPP, aprobación, vigencia y seguimiento del PMD."],
            ["Ley núm. 498-06 de Planificación e Inversión Pública", "Articula planificación, programación institucional e inversión pública."],
            ["Ley núm. 1-12 de la Estrategia Nacional de Desarrollo 2030", "Aporta objetivos nacionales de largo plazo para las políticas territoriales."],
            ["Ley núm. 368-22 de Ordenamiento Territorial, Uso de Suelo y Asentamientos Humanos", "Vincula los instrumentos municipales con el marco nacional de ordenamiento territorial."],
            ["Ley núm. 64-00 de Medio Ambiente y Recursos Naturales", "Sustenta la protección ambiental y la gestión de recursos naturales."],
            ["Ley núm. 147-02 sobre Gestión de Riesgos", "Integra prevención, reducción de riesgos y preparación ante emergencias."],
            ["Ley núm. 225-20 de Gestión Integral y Coprocesamiento de Residuos Sólidos", "Orienta prevención, recolección, aprovechamiento y disposición final de residuos."],
        ],
        [3100, 6260],
    )
    paragraph(
        doc,
        "Conforme a los artículos 122 a 125 de la Ley núm. 176-07, la OMPP coordina la "
        "formulación y evaluación del PMD y el CDM participa en su elaboración, discusión y seguimiento.",
    )
    add_source_note(
        doc,
        "Ley núm. 176-07, arts. 122–125; leyes núm. 498-06, 1-12, 368-22, 64-00, 147-02 y 225-20.",
    )

    add_page_break(doc)
    doc.add_heading("1.7 Metodología", level=2)
    paragraph(
        doc,
        "La elaboración sigue la Guía para la formulación de planes de desarrollo municipales "
        "publicada por el Ministerio de Economía, Planificación y Desarrollo (MEPyD). "
        "La guía organiza el proceso en cuatro etapas enlazadas:",
    )
    methodology = [
        "Organización del proceso. Se conforma el equipo, se acuerdan responsabilidades entre Alcaldía, Concejo de Regidores, OMPP, áreas técnicas y CDM, y se prepara el plan de trabajo.",
        "Diagnóstico municipal. Se reúnen y analizan datos estadísticos, cartografía, documentos municipales y conocimiento local. La evidencia permite describir fortalezas y debilidades para revisión del CDM.",
        "Formulación estratégica. Con la participación del CDM se acuerdan visión, objetivos, resultados, líneas de acción, proyectos e indicadores.",
        "Gestión de los instrumentos. El PMD aprobado se vincula con programación, presupuesto, ejecución, seguimiento, evaluación y rendición de cuentas.",
    ]
    for item in methodology:
        doc.add_paragraph(item, style="List Number")
    add_source_note(
        doc,
        "MEPyD, ¿Cómo elaborar un plan municipal de desarrollo? Guía para la formulación "
        "de planes de desarrollo municipales, pp. 19–49; Ley núm. 176-07, arts. 123–125.",
    )


def write_diagnostic_narrative(
    doc: Document,
    municipality: dict[str, Any],
    data: dict[str, Any],
    national: dict[str, Any],
    historical_context: dict[str, Any],
) -> list[dict[str, str]]:
    name = municipality["municipio"]
    basic = data.get("basic") or {}
    households = data.get("households") or {}
    urban = data.get("urban_rural") or {}
    living = data.get("living") or {}
    education = data.get("education") or {}
    education_level = data.get("education_level") or {}
    education_offer = data.get("education_offer") or {}
    economy = data.get("economy") or {}
    tic = data.get("tic") or {}
    health = data.get("health") or {}
    pyramid = data.get("pyramid") or []
    findings: list[dict[str, str]] = []

    doc.add_heading("2.1 Dinámica demográfica y hogares", level=2)
    if basic:
        age = age_structure(pyramid)
        paragraph(
            doc,
            f"La estructura por edad de {name} se distribuye aproximadamente en "
            f"{fmt_pct(age['youth'])} de población menor de 15 años, {fmt_pct(age['working'])} entre 15 y 64 años "
            f"y {fmt_pct(age['older'])} de 65 años o más. La relación de dependencia demográfica estimada es de "
            f"{fmt_number(age['dependency'], 1)} personas potencialmente dependientes por cada 100 en edad de trabajar. "
            "Estos valores orientan la demanda relativa de cuidados, espacios públicos, movilidad, formación y servicios sociales, "
            "pero deben analizarse con la distribución territorial y la capacidad real de atención. [DASH-PYR]",
        )
        urban_pct = pct(urban.get("urbana"), urban.get("poblacion_total"))
        paragraph(
            doc,
            f"El grado de urbanización es {fmt_pct(urban_pct)} y el tamaño medio de los hogares es "
            f"{fmt_number(households.get('personas_por_hogar'), 2)} personas. "
            "Una mayor concentración urbana puede facilitar cobertura de rutas y equipamientos, mientras que una proporción rural "
            "relevante exige criterios de accesibilidad y costos diferenciados. La cifra no sustituye un análisis por barrio, sección o paraje. [DASH-URB] [DASH-HOG]",
        )
        findings.append(
            {
                "theme": "Población y territorio",
                "finding": f"Variación intercensal de {fmt_pct(basic.get('variacion_pct'))}; urbanización de {fmt_pct(urban_pct)}.",
                "implication": "Ajustar capacidad y distribución territorial de servicios.",
                "source": "DASH-BASE / DASH-URB",
            }
        )
    else:
        paragraph(
            doc,
            "No existe una serie municipal separada en el Dashboard. El diagnóstico demográfico queda pendiente de una "
            "desagregación oficial específica; no se usaron promedios del territorio de origen como sustituto.",
        )

    doc.add_heading("2.2 Vivienda, agua, saneamiento y servicios básicos", level=2)
    if living:
        water = service_share(living, "agua_uso_domestico", "del_acueducto_dentro_de_la_vivienda")
        national_water = national_service_share(national.get("living") or {}, "agua_uso_domestico", "del_acueducto_dentro_de_la_vivienda")
        toilet = service_share(living, "servicios_sanitarios", "inodoro")
        no_toilet = service_share(living, "servicios_sanitarios", "no_tiene")
        bottled = service_share(living, "agua_para_beber", "botellones")
        paragraph(
            doc,
            f"El {fmt_pct(water)} de los hogares declara recibir agua del acueducto dentro de la vivienda, "
            f"frente a {fmt_pct(national_water)} a escala nacional. Para agua de beber, el uso de botellones representa "
            f"{fmt_pct(bottled)}. La comparación sugiere una condición de acceso físico y un patrón de consumo, pero no informa "
            "continuidad, presión, potabilidad ni costo. Esos cuatro aspectos deben verificarse antes de formular intervenciones. [DASH-VIDA]",
        )
        paragraph(
            doc,
            f"En saneamiento, {fmt_pct(toilet)} de los hogares reporta inodoro y {fmt_pct(no_toilet)} declara no disponer "
            "de servicio sanitario. La línea base permite dimensionar la brecha general, aunque no identifica sectores ni calidad "
            "de disposición final. El levantamiento municipal debe localizar hogares críticos y distinguir competencias del ayuntamiento, "
            "INAPA, salud pública y otros actores. [DASH-VIDA]",
        )
        lighting = (living.get("servicios") or {}).get("alumbrado", {})
        grid_share = pct(
            (lighting.get("categorias") or {}).get("energia_eletrica_del_tendido_publico"),
            lighting.get("total"),
        )
        waste = service_share(living, "eliminacion_basura", "la_recoge_el_ayuntamiento")
        burn = service_share(living, "eliminacion_basura", "la_queman")
        propane = service_share(living, "combustible_cocinar", "gas_propano")
        paragraph(
            doc,
            f"La red pública es la fuente principal de alumbrado para {fmt_pct(grid_share)} de los hogares. "
            f"La recogida de residuos por el ayuntamiento alcanza {fmt_pct(waste)}, mientras {fmt_pct(burn)} declara quemarlos. "
            f"El gas propano es el combustible principal de cocina en {fmt_pct(propane)} de los hogares. "
            "La cobertura de recogida no equivale a frecuencia, regularidad ni manejo ambientalmente adecuado; el plan operativo "
            "de residuos debe contrastar rutas, horarios, disposición final y sectores no servidos. [DASH-VIDA]",
        )
        findings.extend(
            [
                {
                    "theme": "Agua y saneamiento",
                    "finding": f"Agua de acueducto dentro: {fmt_pct(water)}; hogares sin sanitario: {fmt_pct(no_toilet)}.",
                    "implication": "Localizar brechas y validar continuidad/calidad.",
                    "source": "DASH-VIDA",
                },
                {
                    "theme": "Residuos sólidos",
                    "finding": f"Recogida municipal declarada: {fmt_pct(waste)}; quema: {fmt_pct(burn)}.",
                    "implication": "Revisar cobertura, frecuencia y disposición final.",
                    "source": "DASH-VIDA",
                },
            ]
        )
    else:
        paragraph(doc, "No se dispone de registros municipales separados de vivienda y servicios básicos en el Dashboard.")

    doc.add_heading("2.3 Educación y capacidades humanas", level=2)
    if education:
        offer = education_offer.get("niveles") or {}
        total_centers = education_offer.get("centros_total")
        total_students = sum(float((offer.get(level) or {}).get("matricula", 0) or 0) for level in offer)
        paragraph(
            doc,
            f"La oferta educativa asociada registra {fmt_number(total_centers)} centros y "
            f"{fmt_number(total_students)} matrículas agregadas. De ellas, "
            f"{fmt_number((offer.get('inicial_primario') or {}).get('matricula'))} corresponden a inicial/primario, "
            f"{fmt_number((offer.get('secundario') or {}).get('matricula'))} a secundario y "
            f"{fmt_number((offer.get('adultos') or {}).get('matricula'))} a educación de adultos. "
            "La cobertura corresponde a la asociación territorial usada por el Dashboard y debe verificarse cuando el distrito educativo "
            "abarque más de un municipio. [DASH-EDU]",
        )
        levels = education_level.get("nivel") or education_level.get("niveles") or {}
        total_levels = sum(float((row or {}).get("total", 0) or 0) for row in levels.values())
        none = float((levels.get("ninguno") or {}).get("total", 0) or 0)
        secondary_total = sum(
            float((levels.get(key) or {}).get("total", 0) or 0)
            for key in ("secundaria", "secundaria_media", "superior", "universitaria_superior", "postgrado")
        )
        paragraph(
            doc,
            f"Entre la población de referencia de tres años o más, {fmt_pct(pct(none, total_levels))} aparece sin nivel de instrucción "
            f"y {fmt_pct(pct(secondary_total, total_levels))} registra secundaria o un nivel superior. "
            "La distribución educativa ayuda a orientar alfabetización, formación técnica, inclusión digital y articulación con empleo, "
            "pero no debe interpretarse como calidad del aprendizaje. [DASH-EDU-NIVEL]",
        )
        efficiency = ((education.get("anuario") or {}).get("eficiencia") or {})
        primary = efficiency.get("primario") or {}
        secondary = efficiency.get("secundario") or {}
        paragraph(
            doc,
            f"Para el distrito educativo asociado, la promoción es {fmt_pct(primary.get('promocion'))} en primaria y "
            f"{fmt_pct(secondary.get('promocion'))} en secundaria; el abandono reportado en secundaria es "
            f"{fmt_pct(secondary.get('abandono'))}. La OMPP debe coordinar la lectura con MINERD y evitar atribuir al gobierno local "
            "resultados que corresponden al sistema educativo. [DASH-EDU]",
        )
        findings.append(
            {
                "theme": "Educación",
                "finding": f"Abandono secundario del distrito asociado: {fmt_pct(secondary.get('abandono'))}.",
                "implication": "Validar alcance territorial y coordinar permanencia escolar.",
                "source": "DASH-EDU",
            }
        )
    else:
        paragraph(doc, "El Dashboard no ofrece una ficha educativa municipal separada para este territorio.")

    doc.add_heading("2.4 Economía, establecimientos y empleo", level=2)
    dee = economy.get("dee_2024") or {}
    if dee:
        total_establishments = dee.get("total_establishments")
        total_employees = dee.get("total_employees")
        bands = dee.get("employment_size_bands") or []
        micro = next((row for row in bands if row.get("size_band") == "micro_1_10"), {})
        micro_est_share = pct(micro.get("establishments"), total_establishments)
        paragraph(
            doc,
            f"El Directorio de Empresas y Establecimientos 2024 identifica {fmt_number(total_establishments)} establecimientos "
            f"y un empleo estimado de {fmt_number(total_employees, 1)}, equivalente a "
            f"{fmt_number(dee.get('avg_employees_per_establishment'), 1)} empleos por establecimiento. "
            f"Los establecimientos de 1 a 10 personas representan {fmt_pct(micro_est_share)} del total. "
            "Esto caracteriza la estructura formal observada por el DEE; no incluye necesariamente toda la economía informal, "
            "el autoempleo ni la producción agropecuaria familiar. [DASH-ECO]",
        )
        sectors = dee.get("sectors") or []
        top_sectors = [
            f"{row.get('label')} ({fmt_number(row.get('employees'), 1)} empleos estimados)"
            for row in sectors[:3]
        ]
        top_specialization = dee.get("top_specialization") or {}
        paragraph(
            doc,
            "Las actividades con mayor empleo estimado en la fuente son "
            + "; ".join(top_sectors)
            + ". "
            + (
                f"La mayor especialización relativa se observa en {top_specialization.get('label')} "
                f"(cociente de localización {fmt_number(top_specialization.get('lq'), 2)}). "
                if top_specialization.get("label")
                else ""
            )
            + "El cociente de localización compara concentración relativa, no productividad ni competitividad. [DASH-ECO]",
        )
        findings.append(
            {
                "theme": "Economía local",
                "finding": f"{fmt_pct(micro_est_share)} de los establecimientos observados son micro (1–10).",
                "implication": "Diseñar apoyo empresarial acorde con la estructura local y confirmar informalidad.",
                "source": "DASH-ECO",
            }
        )
    else:
        paragraph(doc, "No se dispone de una ficha económica municipal separada en el Dashboard.")

    doc.add_heading("2.5 Salud, conectividad y acceso a información", level=2)
    centers = health.get("centros") or []
    internet = ((tic.get("internet") or {}).get("rate_used"))
    computer = ((tic.get("computer") or {}).get("rate_used"))
    if centers or tic:
        by_type = Counter((row.get("tipo_centro") or "Sin clasificación").title() for row in centers)
        type_summary = ", ".join(f"{label}: {count}" for label, count in by_type.most_common(5)) or "sin registros"
        paragraph(
            doc,
            f"El Dashboard registra {len(centers)} establecimientos de salud en el municipio. La composición principal es: "
            f"{type_summary}. Este conteo informa presencia física, no cartera de servicios, personal, horario, capacidad, calidad "
            "ni accesibilidad efectiva. La planificación municipal debe validar esos elementos con el Servicio Nacional de Salud. [DASH-SALUD]",
        )
        paragraph(
            doc,
            f"El uso de internet alcanza {fmt_pct(float(internet) * 100 if internet is not None else None)} y el de computadora "
            f"{fmt_pct(float(computer) * 100 if computer is not None else None)}. "
            "La diferencia entre conectividad y disponibilidad de equipo es relevante para trámites, educación, información pública "
            "y participación digital. El dato requiere desagregación por edad, sexo y zona para diseñar acciones focalizadas. [DASH-TIC]",
        )
        findings.append(
            {
                "theme": "Conectividad",
                "finding": f"Internet: {fmt_pct(float(internet) * 100 if internet is not None else None)}; computadora: {fmt_pct(float(computer) * 100 if computer is not None else None)}.",
                "implication": "Precisar brechas por grupos y territorio.",
                "source": "DASH-TIC",
            }
        )
    else:
        paragraph(doc, "No se dispone de registros municipales separados de salud y conectividad en el Dashboard.")

    doc.add_heading("2.6 Contexto territorial, ambiental y continuidad", level=2)
    if historical_context.get("history"):
        if historical_context.get("document_scope") == "predecessor_district":
            paragraph(
                doc,
                "El plan del antiguo distrito municipal aporta antecedentes históricos y territoriales del ámbito "
                "anterior a la creación del municipio. No se presenta como PMD histórico del municipio actual. "
                "Sus descripciones corresponden a otro período y delimitación, por lo que cualquier reutilización "
                "debe comprobarse con información del nuevo ámbito. [ANT-DM]",
            )
        else:
            paragraph(
                doc,
                "El PMD anterior aporta antecedentes históricos y territoriales útiles para comprender la formación del municipio. "
                "Sin embargo, sus descripciones de infraestructura, ambiente, economía o gestión corresponden a otro período. "
                "Antes de reutilizarlas, la OMPP debe clasificarlas como vigentes, modificadas o superadas y registrar la evidencia actual. [PMD-ANT]",
            )
    else:
        paragraph(
            doc,
            "Las fuentes estadísticas utilizadas no describen con suficiente detalle uso del suelo, riesgos, hidrología, residuos, "
            "movilidad o estado de la infraestructura. Este apartado debe completarse con cartografía oficial, instrumentos ambientales "
            "y verificación técnica municipal; no se agregaron afirmaciones de prensa ni datos no comprobables.",
        )
    return findings


def build_strengths_weaknesses(
    data: dict[str, Any],
    national: dict[str, Any],
) -> tuple[list[str], list[str]]:
    living = data.get("living") or {}
    tic = data.get("tic") or {}
    strengths: list[str] = []
    weaknesses: list[str] = []

    metrics = [
        (
            "Acceso a agua del acueducto dentro de la vivienda",
            service_share(living, "agua_uso_domestico", "del_acueducto_dentro_de_la_vivienda"),
            national_service_share(national.get("living") or {}, "agua_uso_domestico", "del_acueducto_dentro_de_la_vivienda"),
        ),
        (
            "Recogida de residuos por el ayuntamiento",
            service_share(living, "eliminacion_basura", "la_recoge_el_ayuntamiento"),
            national_service_share(national.get("living") or {}, "eliminacion_basura", "la_recoge_el_ayuntamiento"),
        ),
        (
            "Hogares con inodoro",
            service_share(living, "servicios_sanitarios", "inodoro"),
            national_service_share(national.get("living") or {}, "servicios_sanitarios", "inodoro"),
        ),
        (
            "Electricidad de la red pública",
            service_share(living, "alumbrado", "energia_electrica_del_tendido_publico"),
            national_service_share(
                national.get("living") or {},
                "alumbrado",
                "energia_electrica_del_tendido_publico",
            ),
        ),
    ]
    for label, local, country in metrics:
        if local is None or country is None:
            continue
        delta = float(local) - float(country)
        text = (
            f"{label}: {fmt_pct(local)}; promedio nacional: {fmt_pct(country)} "
            f"({delta:+.1f} puntos porcentuales)."
        )
        if delta >= 2:
            strengths.append(text)
        elif delta <= -2:
            weaknesses.append(text)

    internet = ((tic.get("internet") or {}).get("rate_used"))
    national_internet = (((national.get("tic") or {}).get("internet") or {}).get("rate_used"))
    if internet is not None and national_internet is not None:
        local_pct = float(internet) * 100
        country_pct = float(national_internet) * 100
        delta = local_pct - country_pct
        text = (
            f"Uso de internet: {fmt_pct(local_pct)}; promedio nacional: "
            f"{fmt_pct(country_pct)} ({delta:+.1f} puntos porcentuales)."
        )
        if delta >= 2:
            strengths.append(text)
        elif delta <= -2:
            weaknesses.append(text)

    if not strengths:
        strengths.append(
            "Las fuentes comparables disponibles no permiten clasificar una fortaleza cuantitativa "
            "con una diferencia mínima de 2 puntos porcentuales frente al promedio nacional."
        )
    if not weaknesses:
        weaknesses.append(
            "Las fuentes comparables disponibles no permiten clasificar una debilidad cuantitativa "
            "con una diferencia mínima de 2 puntos porcentuales frente al promedio nacional."
        )
    return strengths[:6], weaknesses[:6]


def project_proposals(findings: list[dict[str, str]]) -> list[list[str]]:
    rows: list[list[str]] = []
    mapping = {
        "Agua y saneamiento": (
            "Actualización del diagnóstico operativo de agua y saneamiento",
            "Localizar continuidad, calidad, cobertura y hogares críticos; coordinar responsabilidades sectoriales.",
        ),
        "Residuos sólidos": (
            "Optimización de rutas y control del manejo de residuos",
            "Medir frecuencia, sectores no cubiertos, puntos críticos y disposición final.",
        ),
        "Educación": (
            "Mecanismo local de coordinación para permanencia educativa",
            "Validar datos con el distrito educativo y focalizar factores de abandono dentro de las competencias locales.",
        ),
        "Economía local": (
            "Programa de información y servicios para microempresas",
            "Caracterizar necesidades reales, formalidad, acceso a capacitación y articulación con instituciones competentes.",
        ),
        "Conectividad": (
            "Acceso y alfabetización digital municipal",
            "Identificar grupos con menor acceso y usar espacios municipales existentes antes de definir inversión.",
        ),
        "Población y territorio": (
            "Sistema municipal de información territorial",
            "Desagregar indicadores por sectores y mantener línea base anual con fuente y responsable.",
        ),
    }
    for finding in findings:
        if finding["theme"] not in mapping:
            continue
        title, scope = mapping[finding["theme"]]
        rows.append([title, scope, finding["source"], "Pendiente de validación"])
    if not rows:
        rows.append(
            [
                "Sistema municipal de información territorial",
                "Completar la línea base y localizar brechas antes de formular inversión.",
                "Dashboard / OMPP",
                "Pendiente de validación",
            ]
        )
    return rows[:6]


def add_dashboard_images(doc: Document, paths: list[Path], name: str) -> None:
    doc.add_heading("2. Diagnóstico municipal", level=1)
    paragraph(
        doc,
        f"Las {len(paths)} páginas siguientes reproducen la exportación vigente del Dashboard Territorial. "
        "Presentan la línea base estadística antes de su interpretación narrativa e incorporan los indicadores "
        "demográficos, sociales, económicos, institucionales, ambientales y de inversión disponibles. Los datos "
        "conservan el año, la fuente y el ámbito territorial; cuando la fuente solo publica un total provincial, "
        "la página lo identifica como «Provincia completa» y no lo atribuye al municipio.",
    )
    add_data_table(
        doc,
        ["Bloque", "Contenido", "Lectura esperada"],
        [
            ["Demografía", "Población, mapa, pirámides 2010/2022, hogares y viviendas", "Cambio y composición demográfica"],
            ["Condiciones de vida", "Agua, saneamiento, energía, residuos, educación, salud y conectividad", "Coberturas y brechas observadas"],
            ["Resultados y riesgos", "Seguridad, movilidad, incendios, inclusión, primera infancia y deporte", "Dato municipal o contexto provincial, según la fuente"],
            ["Gestión e inversión", "Inversión territorial, SISMAP, obras viales y licencias", "Magnitud, ejecución y capacidad institucional"],
            ["Comparación", "Municipio, provincia, región y país", "Posición relativa sin convertirla automáticamente en prioridad"],
        ],
        [1000, 4300, 4060],
    )
    add_callout(
        doc,
        "Cómo leer las láminas",
        "Cada lámina es una línea base, no una conclusión automática. La interpretación debe conservar tres preguntas: "
        "qué mide el indicador, qué no mide y qué comprobación territorial hace falta antes de decidir.",
    )
    for index, path in enumerate(paths, 1):
        add_page_break(doc)
        doc.add_picture(str(path), width=Inches(6.5))
        picture_paragraph = doc.paragraphs[-1]
        picture_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        picture_paragraph.paragraph_format.space_after = Pt(0)
        picture_paragraph.paragraph_format.keep_together = True
        if index == len(paths):
            add_source_note(
                doc,
                f"DDPT, Dashboard Territorial, ficha de {name}. Exportación generada el {TODAY.isoformat()}.",
            )


def build_document(
    municipality: dict[str, Any],
    code: str,
    adm2_code: str,
    data: dict[str, Any],
    national: dict[str, Any],
    wikipedia: dict[str, Any],
    historical: list[dict[str, Any]],
    territorial_antecedents: list[dict[str, Any]],
    historical_context: dict[str, Any],
    area_km2: float | None,
    dashboard_pages: list[Path],
    territorial_context: dict[str, Any],
    public_projects: list[dict[str, Any]],
    mapa_meta: dict[str, Any],
    output_path: Path,
) -> None:
    name = municipality["municipio"]
    province = municipality["provincia"]
    region = municipality["region"]
    doc = Document()
    configure_document(doc, name, province, code)

    # Cover
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_before = Pt(72)
    spacer.paragraph_format.space_after = Pt(0)
    kicker = doc.add_paragraph()
    kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
    kicker.paragraph_format.space_after = Pt(18)
    add_run(kicker, "REPÚBLICA DOMINICANA · PLANIFICACIÓN MUNICIPAL", bold=True, color=COLORS["gold"], size=10)
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(10)
    add_run(title, "Plan Municipal de Desarrollo", bold=True, color=COLORS["ink"], size=29)
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(5)
    add_run(subtitle, name, bold=True, color=COLORS["teal"], size=20)
    province_p = doc.add_paragraph()
    province_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    province_p.paragraph_format.space_after = Pt(30)
    add_run(province_p, f"{province} · Región {region}", color=COLORS["muted"], size=12)
    period_p = doc.add_paragraph()
    period_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    period_p.paragraph_format.space_after = Pt(70)
    add_run(period_p, PERIOD, bold=True, color=COLORS["blue"], size=16)
    date_p = doc.add_paragraph()
    date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_p.paragraph_format.space_before = Pt(22)
    add_run(date_p, f"Código geográfico {code or 'por confirmar'}", color=COLORS["muted"], size=9)

    add_page_break(doc)
    doc.add_heading("Contenido", level=1)
    add_data_table(
        doc,
        ["Parte", "Contenido"],
        [
            ["1", "Información general: equipo, introducción, identidad territorial, antecedentes, marco jurídico y metodología."],
            ["2", "Diagnóstico municipal: población, vivienda, servicios, educación, salud, economía y territorio."],
            ["3", "Fortalezas y debilidades basadas en el diagnóstico; situaciones que el CDM deberá priorizar."],
            ["4", "Visión Municipal: espacio de decisión del CDM con un ejemplo de formato."],
            ["5", "Inversión pública: proyectos en ejecución o reprogramados identificados específicamente para el municipio en MapaInversiones."],
            ["6", "Plan de acción: ejemplo de registro para convertir los acuerdos del CDM en acciones verificables."],
        ],
        [900, 8460],
    )
    doc.add_heading("Síntesis ejecutiva", level=1)
    basic = data.get("basic") or {}
    urban = data.get("urban_rural") or {}
    if basic.get("poblacion_total"):
        population = basic.get("poblacion_total")
        variation = basic.get("variacion_pct")
        direction = "aumentó" if float(variation or 0) >= 0 else "disminuyó"
        paragraph(
            doc,
            f"{name}, municipio de la provincia {province}, registró {fmt_number(population)} habitantes "
            f"en 2022. Entre 2010 y 2022 su población {direction} "
            f"{fmt_pct(abs(float(variation)) if variation is not None else None)}. "
            f"El {fmt_pct(pct(urban.get('urbana'), urban.get('poblacion_total')))} reside en área urbana y "
            f"el {fmt_pct(pct(urban.get('rural'), urban.get('poblacion_total')))} en área rural. "
            "Estas proporciones determinan la escala y la distribución territorial de los servicios municipales.",
        )
        living = data.get("living") or {}
        water = service_share(living, "agua_uso_domestico", "del_acueducto_dentro_de_la_vivienda")
        waste = service_share(living, "eliminacion_basura", "la_recoge_el_ayuntamiento")
        burn = service_share(living, "eliminacion_basura", "la_queman")
        internet = ((data.get("tic") or {}).get("internet") or {}).get("rate_used")
        paragraph(
            doc,
            f"El diagnóstico registra {fmt_pct(water)} de hogares con agua del acueducto dentro de la vivienda, "
            f"{fmt_pct(waste)} con recogida municipal de residuos y {fmt_pct(burn)} que los quema. "
            f"El uso de internet alcanza {fmt_pct(float(internet) * 100 if internet is not None else None)}. "
            "Los capítulos siguientes presentan las cifras, sus fuentes y sus límites antes de la priorización del CDM.",
        )
    else:
        paragraph(
            doc,
            f"{name} es uno de los municipios incorporados recientemente al clasificador territorial. "
            "Las fuentes estadísticas usadas por el Dashboard todavía no ofrecen una ficha separada para su ámbito actual. "
            "Por esa razón, la información general utiliza únicamente antecedentes directamente identificados para el territorio "
            "y el diagnóstico no traslada cifras del municipio de origen ni promedios provinciales como sustituto.",
        )

    add_page_break(doc)
    write_general_information(
        doc,
        municipality,
        code,
        data,
        wikipedia,
        historical,
        territorial_antecedents,
        historical_context,
        area_km2,
    )

    add_page_break(doc)
    add_dashboard_images(doc, dashboard_pages, name)
    add_page_break(doc)
    doc.add_heading("2. Diagnóstico municipal: lectura e interpretación", level=1)
    paragraph(
        doc,
        "La lectura siguiente explica qué muestran los indicadores y cuáles son sus límites. "
        "Se diferencia entre dato observado, interpretación técnica y asunto que requiere comprobación municipal.",
    )
    findings = write_diagnostic_narrative(doc, municipality, data, national, historical_context)
    territorial_rows = territorial_indicator_rows(municipality, territorial_context)
    if territorial_rows:
        doc.add_heading("2.7 Indicadores territoriales adicionales", level=2)
        paragraph(
            doc,
            "El Dashboard Territorial amplía la línea base con registros administrativos recientes. "
            "La columna de ámbito es parte del dato: los valores provinciales sirven como contexto y "
            "no describen por sí solos la situación exclusiva del municipio.",
        )
        add_data_table(
            doc,
            ["Indicador", "Valor publicado", "Ámbito", "Período", "Fuente"],
            territorial_rows,
            [2200, 2200, 2020, 1300, 1640],
        )
        add_source_note(
            doc,
            "DDPT, Dashboard Territorial, dataset generado el "
            f"{(territorial_context.get('meta') or {}).get('generated') or TODAY.isoformat()}. "
            "La ausencia de registros no se interpreta automáticamente como cero.",
        )

    doc.add_heading("3. Fortalezas y debilidades basadas en el diagnóstico", level=1)
    paragraph(
        doc,
        "La clasificación siguiente utiliza únicamente indicadores comparables del diagnóstico. "
        "Se considera diferencia relevante una separación de al menos dos puntos porcentuales frente al promedio nacional. "
        "El crecimiento poblacional, el número absoluto de establecimientos y otros datos sin denominador comparable "
        "no se clasifican automáticamente como fortalezas o debilidades.",
    )
    strengths, weaknesses = build_strengths_weaknesses(data, national)
    add_data_table(
        doc,
        ["Fortalezas observadas", "Debilidades observadas"],
        [
            ["\n".join(f"• {item}" for item in strengths), "\n".join(f"• {item}" for item in weaknesses)]
        ],
        [4680, 4680],
    )
    if adm2_code:
        add_source_note(
            doc,
            "ONE, X Censo Nacional de Población y Vivienda 2022; comparación con los agregados nacionales "
            "del Dashboard. La regla de ±2 puntos porcentuales se aplica de manera uniforme.",
        )
    else:
        add_source_note(
            doc,
            "Verificación de cobertura del Dashboard, 2026. No se efectuó una clasificación "
            "cuantitativa sin una ficha municipal separada.",
        )
    doc.add_heading("3.1 Situaciones del diagnóstico para priorización", level=2)
    paragraph(
        doc,
        "Los hallazgos siguientes resumen situaciones verificables. No se les asigna orden de prioridad; "
        "el CDM debe contrastarlos con la experiencia de barrios, comunidades y sectores.",
    )
    add_data_table(
        doc,
        ["Tema", "Situación observada", "Fuente"],
        [
            [row["theme"], row["finding"], row["source"]]
            for row in findings
        ]
        or [["Línea base municipal", "La ficha estadística separada no está disponible.", "OMPP / ONE"]],
        [2200, 5100, 2060],
    )

    doc.add_heading("4. Visión Municipal", level=1)
    paragraph(
        doc,
        "La visión municipal, los objetivos y los resultados expresan acuerdos políticos y sociales. "
        "Deben surgir de la participación del CDM y ser aprobados mediante el procedimiento municipal correspondiente. "
        "La fila siguiente muestra únicamente el formato de registro.",
    )
    doc.add_heading("4.1 Ejemplo de registro de una visión y un objetivo", level=2)
    add_data_table(
        doc,
        ["Carácter", "Visión / tema", "Objetivo / resultado"],
        [
            [
                "EJEMPLO DE FORMATO · NO APROBADO",
                f"“En 2028, {name} avanza como un municipio que mejora de manera equitativa "
                "sus servicios y protege los recursos de su territorio.”",
                "Ejemplo: mejorar un servicio priorizado por el CDM, usando un indicador del diagnóstico "
                "para fijar la línea base y medir el resultado.",
            ],
        ],
        [2260, 3740, 3360],
    )
    add_source_note(
        doc,
        "MEPyD, Guía para la formulación de planes de desarrollo municipales, formulación estratégica, pp. 33–45.",
    )

    add_page_break(doc)
    doc.add_heading("5. Inversión pública", level=1)
    as_of = mapa_meta.get("asOf") or mapa_meta.get("sourceCut") or TODAY.isoformat()
    project_count = len(public_projects)
    programmed = sum(float(project.get("budget") or 0) for project in public_projects)
    executed = sum(float(project.get("executed") or 0) for project in public_projects)
    mapa_source = (
        "MapaInversiones, datos abiertos 2026 y perfiles oficiales de proyectos, "
        "https://mapainversiones.gob.do/. Se incluyen únicamente proyectos con alcance municipal "
        "específico y estado EJECUCIÓN o REPROGRAMAR; los proyectos nacionales o provinciales sin "
        "municipio confirmado no se asignan artificialmente. Los importes se presentan en pesos "
        "dominicanos nominales, exactamente en la unidad publicada por la fuente y sin aplicar "
        "multiplicadores; los valores inusualmente bajos deben confirmarse en el perfil oficial."
    )
    if public_projects:
        paragraph(
            doc,
            f"Con corte al {as_of}, MapaInversiones registra {project_count} proyectos en ejecución o "
            f"reprogramados con identificación territorial específica para {name}. La programación "
            f"presupuestaria 2026 asociada suma {fmt_money(programmed)} y la ejecución registrada "
            f"suma {fmt_money(executed)}. Estos proyectos pertenecen a instituciones públicas y no "
            "constituyen por sí mismos acuerdos, compromisos ni proyectos aprobados por el CDM o el ayuntamiento.",
        )
        add_source_note(doc, mapa_source)
        rows = []
        for project in public_projects:
            source_label = (
                "Ubicación: perfil oficial"
                if project.get("locationSource") == "profile"
                else "Ubicación inferida del nombre; validar"
            )
            start = str(project.get("start") or "")[:4]
            end = str(project.get("end") or "")[:4]
            period = "–".join(value for value in (start, end) if value) or "Sin período"
            state = "En ejecución" if project.get("state") == "EJECUCIÓN" else "Reprogramado"
            rows.append(
                [
                    f"{project.get('name') or 'Proyecto sin nombre'}\n"
                    f"SNIP {project.get('code') or 's/n'} · Ficha {project.get('mapProjectId') or 's/n'}\n"
                    f"{source_label}",
                    project.get("institution") or "No disponible",
                    f"{state}\n{period}",
                    fmt_number(project.get("budget")),
                    fmt_number(project.get("executed")),
                ]
            )
        add_data_table(
            doc,
            [
                "Proyecto",
                "Institución",
                "Estado / período",
                "Asignación 2026 (RD$)",
                "Ejecutado 2026 (RD$)",
            ],
            rows,
            [3540, 1980, 1500, 1170, 1170],
        )
    else:
        paragraph(
            doc,
            f"Con corte al {as_of}, no se identificaron en el dataset utilizado proyectos en ejecución "
            f"o reprogramados con una vinculación municipal específica para {name}. Esta constatación "
            "no excluye proyectos nacionales o provinciales que puedan beneficiar al territorio, ni "
            "sustituye la verificación de la OMPP con las instituciones ejecutoras.",
        )
        add_source_note(doc, mapa_source)

    doc.add_heading("6. Plan de acción a acordar por el CDM", level=1)
    paragraph(
        doc,
        "La fila siguiente sirve únicamente para mostrar cómo convertir un acuerdo futuro del CDM "
        "en una acción con línea base, indicador y decisiones verificables.",
    )
    living = data.get("living") or {}
    waste = service_share(living, "eliminacion_basura", "la_recoge_el_ayuntamiento")
    burn = service_share(living, "eliminacion_basura", "la_queman")
    if waste is not None or burn is not None:
        example_action = "Mejoramiento de la gestión integral de residuos."
        example_baseline = (
            f"Línea base 2022: {fmt_pct(waste)} de hogares con recogida municipal "
            f"y {fmt_pct(burn)} que quema residuos."
        )
    else:
        example_action = "Levantamiento y validación de una línea base municipal priorizada."
        example_baseline = "Línea base pendiente: utilizar una fuente oficial específica del nuevo ámbito municipal."
    add_data_table(
        doc,
        ["Carácter", "Acción", "Línea base e indicador", "Decisiones que debe tomar el CDM"],
        [
            [
                "EJEMPLO DE FORMATO · NO APROBADO",
                example_action,
                example_baseline,
                "Alcance territorial, resultado, meta, plazo, responsables, costo, financiamiento y coordinación institucional.",
            ]
        ],
        [1920, 2320, 2520, 2600],
    )
    add_source_note(
        doc,
        "ONE, Censo 2022 cuando existe ficha municipal; MEPyD, Guía para la formulación "
        "de planes de desarrollo municipales, pp. 33–49.",
    )

    add_page_break(doc)
    doc.add_heading("Fuentes y trazabilidad", level=1)
    source_rows = [
        ["GEO-01", "Clasificador geográfico y GeoJSON municipal", "2026", "Sin paginación", "Alta"],
        [
            "MAPA-IP",
            "MapaInversiones: datos abiertos y perfiles de proyectos",
            str(mapa_meta.get("year") or 2026),
            f"{len(public_projects)} proyectos con alcance municipal específico; corte {mapa_meta.get('asOf') or '2026'}",
            "Alta si perfil oficial; media si nombre del proyecto",
        ],
    ]
    if adm2_code:
        source_rows.extend(
            [
                [
                    "DASH-TERR",
                    "DDPT Dashboard Territorial: indicadores administrativos y comparación territorial",
                    str((territorial_context.get("meta") or {}).get("generated") or TODAY.isoformat())[:10],
                    "Dataset municipal y provincial; ámbito rotulado en cada indicador",
                    "Alta / según fuente oficial enlazada",
                ],
                ["DASH-BASE", "IX y X Censos Nacionales de Población y Vivienda", "2010 / 2022", "Dataset municipal", "Alta"],
                ["DASH-PYR", "IX y X Censos: edad y sexo", "2010 / 2022", "Dataset municipal", "Alta"],
                ["DASH-HOG", "X Censo: hogares y viviendas", "2022", "Dataset municipal", "Alta"],
                ["DASH-URB", "X Censo: población urbana y rural", "2022", "Dataset municipal", "Alta"],
                ["DASH-VIDA", "X Censo: condición de vida y servicios", "2022", "Dataset municipal", "Alta"],
                ["DASH-EDU", "Anuario Estadístico Educativo", "2024", "Distrito asociado", "Media"],
                ["DASH-EDU-NIVEL", "X Censo: nivel de instrucción", "2022", "Dataset municipal", "Alta"],
                ["DASH-ECO", "Directorio de Empresas y Establecimientos", "2024", "Dataset municipal", "Media"],
                ["DASH-SALUD", "Registro de establecimientos SNS mostrado por Dashboard", "s/f", "Dataset municipal", "Media"],
                ["DASH-TIC", "X Censo: tecnologías de información", "2022", "Dataset municipal", "Alta"],
            ]
        )
    else:
        source_rows.append(
            [
                "DASH-GAP",
                "Dashboard de Diagnóstico Territorial",
                "2026",
                "Sin ficha separada para el ámbito municipal actual",
                "Alta",
            ]
        )
    if is_municipal_wikipedia(wikipedia):
        source_rows.append(
            [
                "WIKI-01",
                f"Wikipedia en español: {wikipedia.get('title')}",
                wikipedia.get("retrieved_at", TODAY.isoformat()),
                "Sin paginación",
                "Baja / secundaria",
            ]
        )
    for index, row in enumerate(
        sorted(historical, key=lambda item: item.get("period_end") or 0, reverse=True), 1
    ):
        source_rows.append(
            [
                f"PMD-ANT-{index:02d}",
                f"PMD de {name} {row.get('period_start')}-{row.get('period_end')}",
                f"{row.get('period_start')}-{row.get('period_end')}",
                f"{row.get('pages') or 's/p'} páginas; citas específicas en texto",
                "Alta" if str(row.get("validation_status", "")).startswith("verified") else "Media",
            ]
        )
    for index, row in enumerate(
        sorted(
            territorial_antecedents,
            key=lambda item: item.get("period_end") or 0,
            reverse=True,
        ),
        1,
    ):
        source_rows.append(
            [
                f"ANT-DM-{index:02d}",
                f"Plan del antiguo Distrito Municipal de {name} "
                f"{row.get('period_start')}-{row.get('period_end')}",
                f"{row.get('period_start')}-{row.get('period_end')}",
                f"{row.get('pages') or 's/p'} páginas; antecedente territorial",
                "Alta para el ámbito y período original",
            ]
        )
    add_data_table(doc, ["ID", "Documento", "Año", "Página / cobertura", "Confianza"], source_rows, [1300, 3300, 1300, 2200, 1260])
    add_source_note(
        doc,
        "Los datasets del Dashboard no tienen paginación. Cuando se usa texto de un documento anterior, "
        "el rango de páginas y su alcance territorial se consignan en el párrafo correspondiente.",
    )
    doc.core_properties.title = f"Plan Municipal de Desarrollo de {name} {PERIOD}"
    doc.core_properties.subject = "Información general, diagnóstico territorial e inversión pública municipal con trazabilidad"
    doc.core_properties.author = "DDPT"
    doc.core_properties.keywords = "PMD, municipio, diagnóstico, OMPP, CDM, trazabilidad"
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)


def first_index(rows: Any) -> dict[str, Any]:
    if isinstance(rows, dict):
        return {base.code_key(key): value for key, value in rows.items()}
    result: dict[str, Any] = {}
    for row in rows or []:
        key = base.code_key(row.get("adm2_code"))
        if key and key not in result:
            result[key] = row
    return result


def build_data_bundle(
    adm2_code: str,
    datasets: dict[str, Any],
    adm2_map_2010: dict[str, str | None],
) -> dict[str, Any]:
    if not adm2_code:
        return {
            "basic": None,
            "households": None,
            "urban_rural": None,
            "living": None,
            "education": None,
            "education_level": None,
            "education_offer": None,
            "economy": None,
            "tic": None,
            "health": None,
            "pyramid": [],
            "pyramid2010": [],
            "household_size": [],
        }
    old_code = adm2_map_2010.get(adm2_code)
    return {
        "basic": datasets["basic"].get(adm2_code),
        "households": datasets["households"].get(adm2_code),
        "urban_rural": datasets["urban_rural"].get(adm2_code),
        "living": datasets["living"].get(adm2_code),
        "education": datasets["education"].get(adm2_code),
        "education_level": datasets["education_level"].get(adm2_code),
        "education_offer": datasets["education_offer"].get(adm2_code),
        "economy": datasets["economy"].get(adm2_code),
        "tic": datasets["tic"].get(adm2_code),
        "health": datasets["health"].get(adm2_code),
        "pyramid": (datasets["pyramid"].get(adm2_code) or {}).get("age_groups", []),
        "pyramid2010": datasets["age2010"].get(base.code_key(old_code), []) if old_code else [],
        "household_size": datasets["household_size"].get(adm2_code, []),
    }


def load_datasets(data_dir: Path) -> tuple[dict[str, Any], dict[str, Any], dict[str, str | None]]:
    datasets = {
        "basic": first_index(load_json(data_dir / "indicadores_basicos.json")),
        "households": first_index(load_json(data_dir / "hogares_resumen.json")),
        "urban_rural": first_index(load_json(data_dir / "poblacion_urbana_rural.json")),
        "living": first_index(load_json(data_dir / "condicion_vida.json")),
        "education": first_index(load_json(data_dir / "educacion.json")),
        "education_level": first_index(load_json(data_dir / "educacion_nivel.json")),
        "education_offer": first_index(load_json(data_dir / "educacion_oferta_municipal.json")),
        "economy": first_index(load_json(data_dir / "economia_empleo.json")),
        "tic": first_index(load_json(data_dir / "tic.json")),
        "health": first_index(load_json(data_dir / "salud_establecimientos.json")),
        "pyramid": first_index(load_json(data_dir / "pyramids.json")),
        "age2010": group_index(load_json(data_dir / "edad_sexo_2010.json")),
        "household_size": group_index(load_json(data_dir / "tamano_hogar.json")),
    }
    national = {
        "basic": load_json(data_dir / "national_basic.json"),
        "hogares": load_json(data_dir / "national_hogares.json"),
        "living": load_json(data_dir / "national_condicion_vida.json"),
        "education_level": load_json(data_dir / "national_educacion_nivel.json"),
        "education_offer": load_json(data_dir / "national_educacion_oferta.json"),
        "economy": load_json(data_dir / "national_economia_empleo.json"),
        "tic": load_json(data_dir / "national_tic.json"),
        "health": load_json(data_dir / "national_salud_establecimientos.json"),
    }
    adm2_map_2010 = {
        base.code_key(key): (str(value).zfill(6) if value else None)
        for key, value in load_json(data_dir / "adm2_map_2010.json").items()
    }
    return datasets, national, adm2_map_2010


def main() -> int:
    parser = argparse.ArgumentParser()
    parser.add_argument("--workspace", type=Path, required=True)
    parser.add_argument("--portal-repo", type=Path, required=True)
    parser.add_argument(
        "--mapa-data",
        type=Path,
        required=True,
        help="MapaInversiones municipal dataset produced by IPRegional2",
    )
    parser.add_argument(
        "--dashboard-repo",
        type=Path,
        help="Current DDPT_Dashboard-Territorial repository",
    )
    parser.add_argument(
        "--diagnosticos-dir",
        type=Path,
        help="Directory containing the 158 current Dashboard PDF exports",
    )
    parser.add_argument("--municipality", help="Generate one municipality as a prototype")
    parser.add_argument("--all", action="store_true", help="Generate all 104 targets")
    parser.add_argument("--web-copy", action="store_true", help="Copy DOCX files to the portal downloads folder")
    parser.add_argument("--skip-wikipedia-refresh", action="store_true")
    parser.add_argument("--resume", action="store_true", help="Keep an existing rich DOCX instead of rebuilding it")
    parser.add_argument(
        "--reuse-existing-dashboard-pages",
        action="store_true",
        help="Regenerate the 2010/2022 demographic page but reuse existing pages 2-5",
    )
    parser.add_argument(
        "--reuse-all-dashboard-pages",
        action="store_true",
        help="Reuse all five existing diagnostic images when only the Word narrative changes",
    )
    parser.add_argument(
        "--reuse-dashboard-pdf-pages",
        action="store_true",
        help="Reuse previously rasterized pages from the current Dashboard PDF",
    )
    parser.add_argument("--partition-count", type=int, default=1)
    parser.add_argument("--partition-index", type=int, default=0)
    parser.add_argument(
        "--partial",
        action="store_true",
        help="Generate one partition without replacing the complete manifest",
    )
    args = parser.parse_args()

    workspace = args.workspace.resolve()
    portal_repo = args.portal_repo.resolve()
    dashboard_repo = (
        args.dashboard_repo.resolve()
        if args.dashboard_repo
        else workspace.parent.parent
        / "Plan Regional de Desarrollo"
        / "DDPT_Dashboard-Territorial"
    )
    if not (dashboard_repo / "public" / "data" / "municipios_index.json").exists():
        raise RuntimeError(f"Dashboard repository not found: {dashboard_repo}")
    diagnosticos_dir = (
        args.diagnosticos_dir.resolve()
        if args.diagnosticos_dir
        else portal_repo / "public" / "downloads" / "diagnosticos"
    )
    mapa_payload = load_json(args.mapa_data.resolve())
    mapa_meta = mapa_payload.get("meta", {})
    mapa_projects_by_territory = index_mapa_projects(mapa_payload)
    data_dir = dashboard_repo / "public" / "data"
    output_dir = workspace / "output" / "pmd_borradores_2025_2028"
    docx_dir = output_dir / "docx"
    assets_dir = output_dir / "assets"
    manifest_path = output_dir / "manifest.json"
    wikipedia_cache_path = output_dir / "wikipedia_cache.json"
    historical_cache_path = output_dir / "historical_context_cache.json"
    inventory_path = workspace / "output" / "pmd_historico_fuentes_no_sismap" / "inventory.json"
    source_audit_path = (
        workspace
        / "output"
        / "source_link_audit_2026-07-28"
        / "source_links_162.json"
    )
    classifier_path = workspace / "output" / "clasificador_geografico" / "clasificador_geografico_01-10.txt"
    portal_data_path = portal_repo / "app" / "data" / "municipios.json"
    web_dir = portal_repo / "public" / "downloads" / "pmd-borradores"
    geojson_path = portal_repo / "public" / "data" / "adm2.geojson"

    portal_data = load_json(portal_data_path)
    targets = [
        item
        for item in portal_data
        if not item.get("pmd", {}).get("hasOfficialEvidence")
        and not (item.get("pmd", {}).get("has7_12") or item.get("pmd", {}).get("hasDraft"))
    ]
    if len(targets) != 104:
        raise RuntimeError(f"Expected 104 target municipalities; found {len(targets)}")
    if args.partition_count < 1 or not 0 <= args.partition_index < args.partition_count:
        raise RuntimeError("Invalid partition settings")
    if args.municipality:
        wanted = base.clean(args.municipality)
        targets = [item for item in targets if base.clean(item["municipio"]) == wanted]
        if not targets:
            raise RuntimeError(f"Target municipality not found: {args.municipality}")
    elif not args.all:
        parser.error("Choose --municipality NAME or --all")
    elif args.partition_count > 1:
        targets = [
            item
            for index, item in enumerate(targets)
            if index % args.partition_count == args.partition_index
        ]

    dashboard_index = load_json(data_dir / "municipios_index.json")
    dashboard_by_territory = {
        base.territory_key(item["municipio"], item["provincia"]): base.code_key(item["adm2_code"])
        for item in dashboard_index
    }
    datasets, national, adm2_map_2010 = load_datasets(data_dir)
    territorial_payload = load_json(data_dir / "territorial-dashboard.json")
    territorial_by_code = {
        base.code_key(item.get("code")): item
        for item in territorial_payload.get("municipalities", [])
    }
    territorial_province_by_name = {
        base.clean(item.get("name", "")): item
        for item in territorial_payload.get("provinces", [])
    }
    geojson = load_json(geojson_path)
    geo_by_territory = {
        base.territory_key(
            feature.get("properties", {}).get("municipio", ""),
            feature.get("properties", {}).get("provincia", ""),
        ): feature
        for feature in geojson.get("features", [])
    }

    historical_rows = load_json(inventory_path).get("rows", [])
    historical_by_territory: dict[str, list[dict[str, Any]]] = defaultdict(list)
    antecedents_by_territory: dict[str, list[dict[str, Any]]] = defaultdict(list)
    for row in historical_rows:
        if row.get("status") == "downloaded":
            key = base.territory_key(row["municipality"], row["province"])
            copy = dict(row)
            if key in NEW_MUNICIPALITIES:
                copy["document_scope"] = "predecessor_district"
                antecedents_by_territory[key].append(copy)
            else:
                copy["document_scope"] = "municipality_historical_pmd"
                historical_by_territory[key].append(copy)

    classifier = base.read_classifier_codes(classifier_path)
    wikipedia_cache = load_json(wikipedia_cache_path) if wikipedia_cache_path.exists() else {}
    source_audit = load_json(source_audit_path) if source_audit_path.exists() else {}
    source_audit_by_territory = {
        base.territory_key(row["municipio"], row["provincia"]): row
        for row in source_audit.get("municipalities", [])
    }
    for item in targets:
        key = base.territory_key(item["municipio"], item["provincia"])
        audited = (source_audit_by_territory.get(key) or {}).get("wikipedia") or {}
        if audited.get("status") == "verified_direct" and audited.get("title"):
            previous_wiki = wikipedia_cache.get(key, {})
            same_title = previous_wiki.get("title") == audited.get("title")
            wikipedia_cache[key] = {
                **(
                    previous_wiki
                    if same_title
                    else {
                        "full_extract": "",
                        "sections": {},
                    }
                ),
                "status": "verified",
                "title": audited["title"],
                "url": audited.get("url", ""),
                "audit_method": audited.get("method", ""),
                "has_history": bool(audited.get("has_history")),
                "retrieved_at": TODAY.isoformat(),
            }
        elif audited.get("status") == "not_found":
            wikipedia_cache[key] = {
                "status": "not_found",
                "title": "",
                "url": "",
                "reason": "No se verificó una página municipal directa.",
                "retrieved_at": TODAY.isoformat(),
            }
    if not args.skip_wikipedia_refresh:
        if not source_audit_by_territory:
            base.fetch_wikipedia_exact_batches(targets, wikipedia_cache)
        sanitize_wikipedia_cache(wikipedia_cache, targets)
        fetch_full_wikipedia(wikipedia_cache, targets)
        save_json(wikipedia_cache_path, wikipedia_cache)
    historical_cache = load_json(historical_cache_path) if historical_cache_path.exists() else {}

    previous_manifest = load_json(manifest_path).get("municipalities", []) if manifest_path.exists() else []
    previous_by_id = {row["id"]: row for row in previous_manifest}
    generated_rows: list[dict[str, Any]] = []
    docx_dir.mkdir(parents=True, exist_ok=True)
    assets_dir.mkdir(parents=True, exist_ok=True)
    if args.web_copy:
        web_dir.mkdir(parents=True, exist_ok=True)

    for number, item in enumerate(targets, 1):
        key = base.territory_key(item["municipio"], item["provincia"])
        adm2_code = dashboard_by_territory.get(key, "")
        feature = geo_by_territory.get(key)
        map_adm2_code = (
            base.code_key(feature.get("properties", {}).get("adm2_code"))
            if feature
            else adm2_code
        )
        geographic_code = base.resolve_geographic_code(item, adm2_code, classifier)
        if not geographic_code and previous_by_id.get(item["id"]):
            geographic_code = previous_by_id[item["id"]].get("geographic_code", "")
        data = build_data_bundle(adm2_code, datasets, adm2_map_2010)
        historical = historical_by_territory.get(key, [])
        territorial_antecedents = antecedents_by_territory.get(key, [])
        historical_context = extract_historical_context(
            historical or territorial_antecedents,
            historical_cache,
            key,
        )
        historical_context = dict(historical_context)
        historical_context["document_scope"] = (
            "predecessor_district"
            if not historical and territorial_antecedents
            else "municipality_historical_pmd"
        )
        wikipedia = wikipedia_cache.get(key, {"status": "not_found"})
        public_projects = mapa_projects_by_territory.get(key, [])
        previous = previous_by_id.get(item["id"], {})
        file_name = previous.get("file_name") or (
            f"{geographic_code or f'id-{item['id']:03d}'}_"
            f"PMD_{base.slugify(item['municipio'])}_Borrador_Tecnico_{PERIOD}.docx"
        )
        output_path = docx_dir / file_name
        asset_dir = assets_dir / base.slugify(item["municipio"])
        pdf_matches = sorted(diagnosticos_dir.glob(f"{adm2_code}_*.pdf")) if adm2_code else []
        if adm2_code and len(pdf_matches) != 1:
            raise RuntimeError(
                f"Expected one current Dashboard PDF for {item['municipio']} ({adm2_code}); "
                f"found {len(pdf_matches)}"
            )
        dashboard_pdf_path = pdf_matches[0] if pdf_matches else None
        territorial_context = {
            "municipality": territorial_by_code.get(adm2_code),
            "province": territorial_province_by_name.get(base.clean(item["provincia"])),
            "meta": territorial_payload.get("meta", {}),
            "sources": territorial_payload.get("sources", {}),
        }
        pages: list[Path] = []
        if not (args.resume and output_path.exists()):
            if dashboard_pdf_path:
                pages = rasterize_dashboard_pdf(
                    dashboard_pdf_path,
                    asset_dir,
                    reuse_existing=(
                        args.reuse_dashboard_pdf_pages or args.reuse_all_dashboard_pages
                    ),
                )
            else:
                existing_pages = [asset_dir / f"diagnostico-{index}.png" for index in range(1, 6)]
                if args.reuse_all_dashboard_pages and all(path.exists() for path in existing_pages):
                    pages = existing_pages
                else:
                    pages = generate_dashboard_pages(
                        item,
                        map_adm2_code,
                        data,
                        national,
                        geojson,
                        asset_dir,
                        reuse_existing_rest=args.reuse_existing_dashboard_pages,
                    )
            area_km2 = (feature or {}).get("properties", {}).get("km2")
            build_document(
                item,
                geographic_code,
                adm2_code,
                data,
                national,
                wikipedia,
                historical,
                territorial_antecedents,
                historical_context,
                area_km2,
                pages,
                territorial_context,
                public_projects,
                mapa_meta,
                output_path,
            )
        if args.web_copy:
            shutil.copy2(output_path, web_dir / file_name)
        generated_rows.append(
            {
                "id": item["id"],
                "municipio": item["municipio"],
                "provincia": item["provincia"],
                "region": item["region"],
                "geographic_code": geographic_code,
                "dashboard_adm2_code": adm2_code,
                "dashboard_available": bool(adm2_code),
                "historical_pmd_count": len(historical),
                "historical_periods": [
                    f"{row.get('period_start')}-{row.get('period_end')}" for row in historical
                ],
                "territorial_antecedent_count": len(territorial_antecedents),
                "territorial_antecedent_periods": [
                    f"{row.get('period_start')}-{row.get('period_end')}"
                    for row in territorial_antecedents
                ],
                "wikipedia_status": wikipedia.get("status"),
                "wikipedia_url": wikipedia.get("url", ""),
                "content_version": CONTENT_VERSION,
                "information_general_status": "precompleted",
                "diagnostic_status": "precompleted" if adm2_code else "source-gap",
                "dashboard_dataset_generated": (
                    territorial_payload.get("meta", {}).get("generated")
                ),
                "dashboard_pdf_filename": (
                    dashboard_pdf_path.name if dashboard_pdf_path else ""
                ),
                "dashboard_pdf_pages": (
                    len(pages)
                    if pages
                    else len(list(asset_dir.glob("dashboard-pdf-page-*.jpg")))
                    or len(list(asset_dir.glob("diagnostico-*.png")))
                ),
                "mapa_project_count": len(public_projects),
                "mapa_budget_2026": round(
                    sum(float(project.get("budget") or 0) for project in public_projects),
                    2,
                ),
                "mapa_executed_2026": round(
                    sum(float(project.get("executed") or 0) for project in public_projects),
                    2,
                ),
                "mapa_as_of": mapa_meta.get("asOf") or mapa_meta.get("sourceCut"),
                "file_name": file_name,
                "relative_url": f"downloads/pmd-borradores/{file_name}",
            }
        )
        print(f"[{number}/{len(targets)}] {item['municipio']} -> {output_path.name}")

    if args.partial:
        print(
            json.dumps(
                {
                    "generated": len(generated_rows),
                    "partition_index": args.partition_index,
                    "partition_count": args.partition_count,
                    "partial": True,
                },
                ensure_ascii=False,
            )
        )
        return 0

    save_json(historical_cache_path, historical_cache)
    if args.municipality:
        generated_ids = {row["id"] for row in generated_rows}
        combined = [row for row in previous_manifest if row["id"] not in generated_ids] + generated_rows
    else:
        combined = generated_rows
    combined = sorted(combined, key=lambda row: row["id"])
    manifest = {
        "generated_at": TODAY.isoformat(),
        "period": PERIOD,
        "content_version": CONTENT_VERSION,
        "target_rule": "No PMD official evidence and no existing 7-12/draft",
        "expected_target_count": 104,
        "generated_count": len(combined),
        "municipalities": combined,
    }
    save_json(manifest_path, manifest)
    save_json(portal_repo / "public" / "data" / "generated-pmd-drafts.json", manifest)

    csv_path = output_dir / "manifest.csv"
    columns = [
        "id",
        "municipio",
        "provincia",
        "region",
        "geographic_code",
        "dashboard_adm2_code",
        "dashboard_available",
        "historical_pmd_count",
        "historical_periods",
        "territorial_antecedent_count",
        "territorial_antecedent_periods",
        "wikipedia_status",
        "wikipedia_url",
        "content_version",
        "information_general_status",
        "diagnostic_status",
        "dashboard_dataset_generated",
        "dashboard_pdf_filename",
        "dashboard_pdf_pages",
        "mapa_project_count",
        "mapa_budget_2026",
        "mapa_executed_2026",
        "mapa_as_of",
        "file_name",
        "relative_url",
    ]
    with csv_path.open("w", encoding="utf-8-sig", newline="") as handle:
        writer = csv.DictWriter(handle, fieldnames=columns)
        writer.writeheader()
        for row in combined:
            copy = dict(row)
            copy["historical_periods"] = "; ".join(copy.get("historical_periods", []))
            copy["territorial_antecedent_periods"] = "; ".join(
                copy.get("territorial_antecedent_periods", [])
            )
            writer.writerow(copy)
    print(
        json.dumps(
            {
                "generated": len(generated_rows),
                "manifest_total": len(combined),
                "output_dir": str(output_dir),
                "web_copy": args.web_copy,
            },
            ensure_ascii=False,
        )
    )
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
